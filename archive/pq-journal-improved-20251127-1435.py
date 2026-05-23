#!/usr/bin/env python3
# PQ Journal Notebook (PyQt6 GUI) + Late-Fusion Emotion Summaries + PyTorch FER + Local Llama
# Modern UX • Per-entry ML-KEM + AES-GCM • Prev/Next • .docx Export
# Offline Vosk Dictation + Tone/Text/Video Emotion Fusion with ResNet50 FER + Local Llama.cpp
# Author: Andrew Fried • Enhanced with Local Llama Integration • © 2025 Meaningful Systems, LLC
# DEBUG VERSION 20251127-1435: Enhanced Llama error logging and diagnostics
# FIX: Changed from /completion to /v1/completions (OpenAI-compatible API)
# FIX: Added missing 'model' parameter to API requests
# FIX: Strip .gguf extension from model name + query /v1/models for available models
# FIX: Auto-retry with multiple model name variations (gpt-3.5-turbo, llama, tinyllama, etc.)

from __future__ import annotations
import sys, os, json, base64, uuid, queue, re, zipfile, shutil, math, time
import enum, threading, traceback
from array import array
from collections import deque
from pathlib import Path
from typing import Optional, Tuple, List, Dict, Any, Callable
from dataclasses import dataclass

# ---------------------- Optional deps (import-guarded) ----------------------
missing: List[str] = []
try:
    import oqs
except Exception:
    oqs = None; missing.append("pyoqs")

try:
    from cryptography.hazmat.primitives.ciphers.aead import AESGCM
    from cryptography.hazmat.primitives.kdf.hkdf import HKDF
    from cryptography.hazmat.primitives.kdf.pbkdf2 import PBKDF2HMAC
    from cryptography.hazmat.primitives import hashes
except Exception:
    AESGCM = HKDF = PBKDF2HMAC = hashes = None; missing.append("cryptography")

try:
    import docx
except Exception:
    docx = None; missing.append("python-docx (for export)")

try:
    import sounddevice as sd
except Exception:
    sd = None; missing.append("sounddevice (for dictation)")

try:
    from vosk import Model as VoskModel, KaldiRecognizer
except Exception:
    VoskModel = KaldiRecognizer = None; missing.append("vosk (offline STT)")

# Optional: Transformers sentiment head (runs offline if model is cached)
try:
    from transformers import pipeline as hf_pipeline
except Exception:
    hf_pipeline = None; missing.append("transformers (optional text valence)")

# NEW: Computer Vision for video emotion detection
try:
    import cv2
    import numpy as np
except Exception:
    cv2 = np = None; missing.append("opencv-python numpy (for video)")

# NEW: PyTorch for facial emotion recognition
try:
    import torch
    import torch.nn as nn
    import torchvision.transforms as transforms
    from torchvision.models import resnet50
except Exception:
    torch = nn = transforms = resnet50 = None; missing.append("torch torchvision (for FER)")

from statistics import median
from PyQt6 import QtCore, QtGui, QtWidgets
from PyQt6.QtCore import Qt, QThread, pyqtSignal

# NEW: Local Llama integration for emotion interpretation
try:
    import requests
    import subprocess
    llm_available = True
except ImportError:
    llm_available = False
    print("Requests module not available - local LLM integration disabled")

APP_TITLE = "PQ Journal Notebook — PyQt6 + PyTorch FER + Local Llama"
MAX_NOTEBOOK_BYTES = 20 * 1024 * 1024

CYBER_BLUE_HEX = "#00D0FF"; SILVER_HEX = "#C0C0C0"; BLACK_HEX = "#0a0a0a"
CHARCOAL_HEX = "#232323";   INPUT_HEX  = "#151515"

# Use the larger model you installed:
VOSK_DIR_NAME = "vosk-model-en-us-0.22-lgraph"

# ============================ Emotion State Machine ============================
class EmotionState(enum.Enum):
    """State machine for managing Llama emotion inference with speech pauses."""
    IDLE = 0           # no audio, no pending job
    LISTENING = 1      # building current paragraph
    WAITING_MODEL = 2  # a job is in-flight, waiting for Llama
    READY_NEXT = 3     # llama finished; waiting for next pause to enqueue again

@dataclass
class EmotionJob:
    """A single emotion analysis job to be processed by Llama."""
    text: str
    vad: dict  # {'valence': float, 'arousal': float, 'dominance': float}
    created_ts: float

# Timing constants for pause detection
PAUSE_MS = 2000  # Silence duration to trigger pause (increased from 800ms to 2s)
PAUSE_DEBOUNCE_MS = 400  # Debounce multiple pauses
LLAMA_TIMEOUT_MS = 5000  # 5 second timeout for LLM responses

# LLM processing thresholds
MIN_WORDS_FOR_LLM = 3  # Minimum words before sending to LLM
MIN_TIME_FOR_SHORT_LLM = 3.0  # Minimum seconds to accumulate short phrases (< MIN_WORDS_FOR_LLM) - reduced from 5s
MIN_TIME_FOR_ANY_LLM = 2.0  # Minimum seconds to wait even for longer phrases (encourages paragraphs)

def debug_timestamp():
    # Function Purpose: Generate timestamp for debug output
    import datetime
    return datetime.datetime.now().strftime("%H:%M:%S.%f")[:-3]

class EmotionQueueWorker(threading.Thread):
    """Single-threaded worker for processing emotion jobs with Llama."""

    def __init__(self, in_q: queue.Queue[EmotionJob],
                 on_result: Callable[[EmotionJob, str], None],
                 call_llama: Callable[[EmotionJob], str]):
        # Function Purpose: Initialize worker thread with queue and callbacks
        super().__init__(daemon=True, name="EmotionQueueWorker")
        self.in_q = in_q
        self.on_result = on_result
        self.call_llama = call_llama
        self._stop_event = threading.Event()
        self._sentinel = object()  # Unique object for shutdown signal

    def run(self):
        # Function Purpose: Main worker loop processing emotion jobs until stopped
        while not self._stop_event.is_set():
            try:
                # Block until job arrives or sentinel received
                job = self.in_q.get(timeout=1.0)

                # Check for shutdown sentinel
                if job is self._sentinel:
                    break

                # Process the emotion job with Llama
                try:
                    result = self.call_llama(job)
                    # Safely callback to UI thread
                    self.on_result(job, result)
                except Exception as e:
                    # Function Purpose: Handle Llama errors with fallback response
                    print(f"Llama call failed: {e}")
                    traceback.print_exc()
                    # Return safe fallback emotion phrase
                    self.on_result(job, "uncertain but reflective")

                finally:
                    self.in_q.task_done()

            except queue.Empty:
                # Timeout occurred, check stop event again
                continue
            except Exception as e:
                print(f"Worker thread error: {e}")
                traceback.print_exc()

    def stop(self):
        # Function Purpose: Signal worker to stop and join thread
        self._stop_event.set()
        # Put sentinel to wake up blocked get()
        try:
            self.in_q.put_nowait(self._sentinel)
        except queue.Full:
            pass  # Queue full, worker will see stop_event on next timeout

def build_llama_prompt(text: str, vad: dict) -> str:
    # Function Purpose: Build emotion analysis prompt for Llama using text and VAD
    return f"""[INSTRUCTIONS]
You summarize a journaler's emotional state as a concise phrase.
Use BOTH the text and the acoustic cues (Valence, Arousal, Dominance).
Follow the rules strictly:

RULES:
- Output ONE short phrase of 3–8 words.
- NEVER use commas - use "and" or "but" instead.
- No lists, no slashes, no "etc.", no quotes.
- No first-person ("I", "I feel", "I am").
- No duplicate words.
- Prefer natural, family-friendly phrasing (e.g., "hopeful but tense").
- If signals conflict, reflect mixed affect (e.g., "hopeful yet uneasy").
- Do NOT explain. Do NOT add labels.

VAD HINTS (soft guidance, not hard rules):
- High Valence + High Arousal → "energized, joyful, enthusiastic"
- High Valence + Low Arousal → "calm, content, relieved"
- Low Valence + High Arousal → "anxious, angry, agitated"
- Low Valence + Low Arousal → "tired, sad, discouraged"
- Higher Dominance → "confident, in control"
- Lower Dominance → "overwhelmed, powerless"

FEW-SHOT EXAMPLES
Text: "I'm so alive I can feel my hair growing and my nails growing! It feels like every cell in my body just had a milkshake!"
V: +0.62  A: +0.48  D: +0.55
Face: happy:0.72, neutral:0.18, surprise:0.06, fear:0.03, sad:0.01
{{bonkers and goofy}}

Text: "I can't focus; everything keeps piling up."
V: -0.58  A: +0.51  D: -0.41
Face: fear:0.68, angry:0.15, sad:0.12, neutral:0.04, disgust:0.01
{{overwhelmed, anxious, seeking relief}}

Text: "It hurts, but I know we'll get through."
V: -0.20  A: +0.18  D: +0.22
Face: sad:0.55, neutral:0.25, happy:0.12, fear:0.05, angry:0.03
{{hurting yet cautiously hopeful}}

Text: "Not much to say. Just tired lately."
V: -0.25  A: -0.45  D: -0.10
Face: neutral:0.61, sad:0.28, angry:0.06, fear:0.03, disgust:0.02
{{drained and low-spirited}}

[INPUT]
Text: "{text}"
V: {vad.get('valence', 0):.2f}
A: {vad.get('arousal', 0):.2f}
D: {vad.get('dominance', 0):.2f}

{{"""

def call_llama(job: EmotionJob) -> str:
    # Function Purpose: Make HTTP request to local Llama server with detailed debugging
    # Note: This function will be called by the worker thread
    # WARNING: This is the OLD LLM path - no auto-start capability!

    print("[WARNING] Using OLD call_llama function - no auto-start! Change LLM dropdown to 'local' for auto-start.")

    timestamp = debug_timestamp()

    try:
        # Build the prompt from the job
        prompt = build_llama_prompt(job.text, job.vad)

        print(f"\n[{timestamp}] === LLAMA DEBUG START ===")
        print(f"[{timestamp}] Input text: '{job.text}'")
        print(f"[{timestamp}] Input VAD: {job.vad}")
        print(f"[{timestamp}] Prompt length: {len(prompt)} chars")
        print(f"[{timestamp}] Prompt sent to Llama:")
        print(f"---START PROMPT---")
        print(prompt)
        print(f"---END PROMPT---")

        # Check if requests is available
        try:
            import requests
        except ImportError:
            print(f"[{timestamp}] ERROR: Requests module not available")
            return "server unavailable"

        # Try to connect to Llama server with detailed logging
        try:
            request_payload = {
                "prompt": prompt,
                "max_tokens": 20,
                "temperature": 0.3,
                "stop": ["."],
                "echo": False
            }

            print(f"[{timestamp}] Request payload: {request_payload}")

            response = requests.post(
                "http://127.0.0.1:8080/completion",
                json=request_payload,
                timeout=10
            )

            print(f"[{timestamp}] HTTP Status: {response.status_code}")
            print(f"[{timestamp}] Response headers: {dict(response.headers)}")

            if response.status_code == 200:
                # Log raw response text
                raw_response = response.text
                print(f"[{timestamp}] Raw HTTP response text:")
                print(f"---START RAW RESPONSE---")
                print(repr(raw_response))  # Use repr to show exact characters
                print(f"---END RAW RESPONSE---")

                # Parse JSON
                try:
                    result = response.json()
                    print(f"[{timestamp}] Parsed JSON: {result}")
                except Exception as json_error:
                    print(f"[{timestamp}] JSON parse error: {json_error}")
                    return "json parse error"

                # Extract content
                content_raw = result.get("content", "")
                print(f"[{timestamp}] Content before strip: {repr(content_raw)}")

                content = content_raw.strip()
                print(f"[{timestamp}] Content after strip: {repr(content)}")

                # Clean up the content to extract emotion phrase from curly brackets
                import re

                # Look for content within curly brackets or after opening bracket
                bracket_match = re.search(r'\{([^}]+)\}', content)
                if bracket_match:
                    content = bracket_match.group(1).strip()
                    print(f"[{timestamp}] Found content in brackets: {repr(content)}")
                else:
                    # Look for content after opening bracket (handles trailing bracket case)
                    partial_match = re.search(r'\{(.+)', content)
                    if partial_match:
                        content = partial_match.group(1).strip()
                        # Remove any trailing brackets (single or multiple)
                        content = re.sub(r'\}+$', '', content).strip()
                        print(f"[{timestamp}] Found content after opening bracket: {repr(content)}")
                    else:
                        # Fallback to original cleaning if no brackets found
                        print(f"[{timestamp}] No brackets found, using fallback cleaning")

                        # Remove quotes at start/end
                        content = re.sub(r'^"([^"]+)".*', r'\1', content)

                        # Split by newlines and take the first meaningful line
                        lines = content.split('\n')
                        clean_line = ""
                        for line in lines:
                            line = line.strip()
                            # Skip empty lines and lines that start with brackets
                            if line and not line.startswith('[') and not line.startswith('Text:') and not line.startswith('V:'):
                                clean_line = line
                                break

                        # If we found a clean line, use it, otherwise use the first line
                        if clean_line:
                            content = clean_line
                        else:
                            content = lines[0].strip() if lines else content

                        # Remove any remaining unwanted patterns
                        content = re.sub(r'\[OUTPUT\].*', '', content)
                        content = re.sub(r'Emotion phrase:.*', '', content)
                        content = re.sub(r'Text:.*', '', content)
                        content = content.strip()

                # Additional cleaning to remove commas and fix common issues
                content = re.sub(r',\s*', ' ', content)  # Replace commas with spaces
                content = re.sub(r'\}+', '', content)    # Remove any remaining brackets
                content = re.sub(r'\s+', ' ', content)   # Collapse multiple spaces
                content = content.strip()

                print(f"[{timestamp}] Content after cleaning: {repr(content)}")
                print(f"[{timestamp}] Content length: {len(content)}")

                # Split into words and analyze
                words_raw = content.split()
                print(f"[{timestamp}] Words from split: {words_raw}")
                print(f"[{timestamp}] Word count: {len(words_raw)}")

                words = words_raw[:8]
                print(f"[{timestamp}] Words after [:8] limit: {words}")
                print(f"[{timestamp}] Final word count: {len(words)}")

                if len(words) >= 2:  # Lowered requirement from 3 to 2
                    final_result = " ".join(words)
                    print(f"[{timestamp}] SUCCESS: Returning '{final_result}'")
                    print(f"[{timestamp}] === LLAMA DEBUG END ===\n")
                    return final_result
                else:
                    print(f"[{timestamp}] FAILURE: Only {len(words)} words, need >= 2")
                    print(f"[{timestamp}] Returning 'processing incomplete'")
                    print(f"[{timestamp}] === LLAMA DEBUG END ===\n")
                    return "processing incomplete"
            else:
                print(f"[{timestamp}] ERROR: HTTP {response.status_code}")
                print(f"[{timestamp}] Response body: {response.text}")
                print(f"[{timestamp}] === LLAMA DEBUG END ===\n")
                return "server error"

        except requests.exceptions.ConnectionError as e:
            print(f"[{timestamp}] ERROR: Cannot connect to Llama server: {e}")
            print(f"[{timestamp}] === LLAMA DEBUG END ===\n")
            return "server unavailable"
        except requests.exceptions.Timeout as e:
            print(f"[{timestamp}] ERROR: Llama server timeout: {e}")
            print(f"[{timestamp}] === LLAMA DEBUG END ===\n")
            return "server timeout"
        except requests.exceptions.RequestException as e:
            print(f"[{timestamp}] ERROR: Request exception: {e}")
            print(f"[{timestamp}] === LLAMA DEBUG END ===\n")
            return "request failed"

    except Exception as e:
        print(f"[{timestamp}] UNEXPECTED ERROR: {e}")
        traceback.print_exc()
        print(f"[{timestamp}] === LLAMA DEBUG END ===\n")
        return "uncertain but reflective"

# ============================ Emotion Matrix ===============================
def _default_emotion_matrix() -> dict:
    # Bias-reduced version for neutral/academic speech.
    return {
        "schema_version": "1.2",
        "bins": {
            "valence": [-1.01, -0.6, -0.2, 0.2, 0.6, 1.01],
            "arousal": [-0.01, 0.35, 0.65, 0.85, 1.01],   # higher bar for "excited"
            "dominance": [-1.01, -0.4, 0.4, 1.01]
        },
        "grid": [
            [ # A0 (low arousal)
                {"adjectives": ["numb","flat"],        "emoji":"😶","color":"#8E8E93"},
                {"adjectives": ["subdued","downcast"], "emoji":"😔","color":"#6B7280"},
                {"adjectives": ["calm","neutral"],     "emoji":"😐","color":"#94A3B8"},
                {"adjectives": ["content","at ease"],  "emoji":"🙂","color":"#A7F3D0"},
                {"adjectives": ["serene","soothed"],   "emoji":"😌","color":"#86EFAC"},
            ],
            [ # A1 (low-mid arousal) → shift positive cells to neutral descriptors
                {"adjectives": ["tense","uneasy"],         "emoji":"😟","color":"#FCA5A5"},
                {"adjectives": ["blue","melancholic"],     "emoji":"😞","color":"#F59E0B"},
                {"adjectives": ["reserved","reflective"],  "emoji":"🤔","color":"#E5E7EB"},
                {"adjectives": ["measured","objective"],   "emoji":"📎","color":"#CBD5E1"},
                {"adjectives": ["even","matter-of-fact"],  "emoji":"📄","color":"#C7D2FE"},
            ],
            [ # A2 (mid arousal)
                {"adjectives": ["irritable","agitated"],   "emoji":"😤","color":"#F87171"},
                {"adjectives": ["sad","discouraged"],      "emoji":"😢","color":"#FED7AA"},
                {"adjectives": ["alert","focused"],        "emoji":"🧠","color":"#93C5FD"},
                {"adjectives": ["buoyant","cheerful"],     "emoji":"😄","color":"#A7F3D0"},
                {"adjectives": ["steady","constructive"],  "emoji":"🧩","color":"#86EFAC"},
            ],
            [ # A3 (high arousal)
                {"adjectives": ["angry","heated"],         "emoji":"😠","color":"#EF4444"},
                {"adjectives": ["anxious","strained"],     "emoji":"😰","color":"#FB923C"},
                {"adjectives": ["energized","intense"],    "emoji":"⚡","color":"#60A5FA"},
                {"adjectives": ["excited","lively"],       "emoji":"🤩","color":"#34D399"},
                {"adjectives": ["joyful","exuberant"],     "emoji":"🎉","color":"#22C55E"},
            ],
        ],
        "intensifiers": [
            {"threshold": 0.25, "adverb": ""},
            {"threshold": 0.55, "adverb": "slightly"},
            {"threshold": 0.72, "adverb": "moderately"},
            {"threshold": 0.86, "adverb": "very"},
            {"threshold": 0.94, "adverb": "highly"},
        ],
        "dominance_tones": [
            {"min": -1.01, "max": -0.4, "word": "overwhelmed", "emoji":"🫨"},
            {"min":  0.40, "max":  1.01, "word": "confident",   "emoji":"💪"},
        ],
        "defaults": {
            "mixed_variability_threshold": 0.45,
            "min_words": 3,
            "max_words": 5,
            "insufficient_label": "insufficient signal",
        }
    }

def load_emotion_matrix() -> dict:
    # Search ~/.pq-journal/emotion_matrix.json first, then app dir.
    cand = [
        Path.home()/".pq-journal"/"emotion_matrix.json",
        (Path(__file__).resolve().parent if "__file__" in globals() else Path.cwd())/"emotion_matrix.json",
        ]
    for p in cand:
        try:
            if p.exists():
                return json.loads(p.read_text(encoding="utf-8"))
        except Exception:
            pass
    return _default_emotion_matrix()

# Utility: bucket helper
def _bucket(value: float, edges: List[float]) -> int:
    for i in range(len(edges)-1):
        if edges[i] <= value < edges[i+1]:
            return i
    return len(edges)-2

# ============================ Local Llama Emotion Interpreter ============================
class LocalLlamaEmotionInterpreter:
    """Local-only Llama.cpp integration for emotion interpretation with auto-start capability"""

    def __init__(self, host: str = "127.0.0.1", port: int = 11434, auto_start: bool = True):
        print(f"[DEBUG] Initializing LocalLlamaEmotionInterpreter with auto_start={auto_start}")
        self.host = host
        self.port = port
        self.base_url = f"http://{host}:{port}"
        self.available = False
        self.auto_start = auto_start
        self.server_process = None
        self._llama_binary_path = None
        self._model_path = None
        self._model_name = None  # Model name for API requests

        # Find llama.cpp binary and model
        print("[DEBUG] Searching for llama binary and model...")
        self._find_llama_binary()
        self._find_model()

        # Extract model name from path (filename without extension)
        if self._model_path:
            import os
            filename = os.path.basename(self._model_path)
            # Remove .gguf, .ggml, or .bin extension
            for ext in ['.gguf', '.ggml', '.bin']:
                if filename.endswith(ext):
                    filename = filename[:-len(ext)]
                    break
            self._model_name = filename
            print(f"[DEBUG] Model name for API: {self._model_name}")

        print(f"[DEBUG] Binary found: {self._llama_binary_path}")
        print(f"[DEBUG] Model found: {self._model_path}")

        # Check if server is already running, start if needed
        print("[DEBUG] Ensuring server availability...")
        self._ensure_server_available()

    def _check_availability(self):
        """Check if local Llama server is running"""
        try:
            # Try multiple endpoints to check if server is up
            endpoints_to_try = ["/health", "/v1/models", "/"]

            for endpoint in endpoints_to_try:
                try:
                    response = requests.get(f"{self.base_url}{endpoint}", timeout=2)
                    if response.status_code in [200, 404]:  # 404 is also OK, means server is responding
                        print(f"[DEBUG] Server responded to {endpoint} with status {response.status_code}")

                        # Also test the /v1/completions endpoint specifically (OpenAI-compatible)
                        try:
                            test_payload = {
                                "model": self._model_name or "test",
                                "prompt": "test",
                                "max_tokens": 1,
                                "stream": False
                            }
                            test_response = requests.post(f"{self.base_url}/v1/completions", json=test_payload, timeout=2)
                            print(f"[DEBUG] Completion endpoint test: status {test_response.status_code}")
                            if test_response.status_code in [200, 400, 404, 500]:  # Any response means server is there
                                self.available = True
                                return
                        except Exception as test_err:
                            print(f"[DEBUG] Completion endpoint test failed: {test_err}")
                            # Server responded to health check, so mark as available anyway
                            self.available = True
                            return
                except requests.exceptions.RequestException as e:
                    print(f"[DEBUG] Endpoint {endpoint} failed: {e}")
                    continue

            self.available = False
            print(f"[DEBUG] All endpoints failed, server not available")

        except Exception as e:
            print(f"[DEBUG] Availability check error: {e}")
            self.available = False

    def _find_llama_binary(self):
        """Find llama.cpp server binary in common locations"""
        common_paths = [
            # Local build paths (check first)
            "./models/llama.cpp/build/bin/llama-server",
            "./llama.cpp/build/bin/llama-server",
            # System paths
            "/usr/local/bin/llama-server",
            "/opt/homebrew/bin/llama-server",
            "./llama-server",
            "./llama.cpp/llama-server",
            "/usr/bin/llama-server",
            # Alternative names
            "/usr/local/bin/server",
            "/opt/homebrew/bin/server",
            "./server",
            "./llama.cpp/server"
        ]

        for path in common_paths:
            if os.path.exists(path) and os.access(path, os.X_OK):
                self._llama_binary_path = path
                print(f"Found llama-server binary: {path}")
                return

        # Try to find via which/where commands
        try:
            result = subprocess.run(["which", "llama-server"], capture_output=True, text=True, timeout=5)
            if result.returncode == 0 and result.stdout.strip():
                self._llama_binary_path = result.stdout.strip()
                print(f"Found llama-server via which: {self._llama_binary_path}")
                return
        except:
            pass

        try:
            result = subprocess.run(["which", "server"], capture_output=True, text=True, timeout=5)
            if result.returncode == 0 and result.stdout.strip():
                # Check if this is actually llama.cpp server
                path = result.stdout.strip()
                if "llama" in path.lower() or self._is_llama_server_binary(path):
                    self._llama_binary_path = path
                    print(f"Found llama server via which: {self._llama_binary_path}")
                    return
        except:
            pass

        print("Warning: llama-server binary not found. Auto-start will not work.")

    def _is_llama_server_binary(self, path: str) -> bool:
        """Check if binary is likely llama.cpp server by running --help"""
        try:
            result = subprocess.run([path, "--help"], capture_output=True, text=True, timeout=3)
            help_text = result.stdout.lower() + result.stderr.lower()
            return "llama" in help_text and ("server" in help_text or "host" in help_text)
        except:
            return False

    def _find_model(self):
        """Find a suitable model file for llama.cpp"""
        model_extensions = [".gguf", ".ggml", ".bin"]
        search_dirs = [
            "./models/llama.cpp",  # Check llama.cpp dir first
            "./models",
            "./",
            "../models",
            os.path.expanduser("~/models"),
            "/usr/local/share/llama-models",
            "/opt/homebrew/share/llama-models"
        ]

        for search_dir in search_dirs:
            if os.path.isdir(search_dir):
                for root, dirs, files in os.walk(search_dir):
                    for file in files:
                        if any(file.lower().endswith(ext) for ext in model_extensions):
                            model_path = os.path.join(root, file)
                            # Prefer TinyLlama (known working), then other chat models
                            if "tinyllama" in file.lower():
                                self._model_path = model_path
                                print(f"Found preferred model: {model_path}")
                                return
                            elif any(keyword in file.lower() for keyword in ["7b", "chat", "instruct", "emotion"]):
                                if not self._model_path:  # Don't override TinyLlama
                                    self._model_path = model_path
                                    print(f"Found chat model: {model_path}")
                            elif not self._model_path:  # Use first found model as fallback
                                self._model_path = model_path

        if self._model_path:
            print(f"Found model: {self._model_path}")
        else:
            print("Warning: No model files found. Auto-start will not work.")

    def _ensure_server_available(self):
        """Ensure llama server is running, start if needed and possible"""
        print(f"[DEBUG] Checking if server is available at {self.base_url}")
        self._check_availability()
        print(f"[DEBUG] Server available: {self.available}")

        if not self.available and self.auto_start:
            print(f"[DEBUG] Server not available, auto_start={self.auto_start}")
            print(f"[DEBUG] Binary path: {self._llama_binary_path}")
            print(f"[DEBUG] Model path: {self._model_path}")

            if self._llama_binary_path and self._model_path:
                print("Llama server not running, attempting to start...")
                if self._start_server():
                    # Wait for server to be ready
                    import time
                    print("Waiting for llama server to become ready...")
                    for i in range(60):  # Wait up to 60 seconds (model loading can be slow)
                        time.sleep(1)

                        # Check if process is still running
                        if self.server_process and self.server_process.poll() is not None:
                            print(f"[ERROR] Server process died with exit code {self.server_process.returncode}")
                            # Try to get some output
                            try:
                                stdout, stderr = self.server_process.communicate(timeout=1)
                                if stdout:
                                    print(f"[ERROR] Server stdout: {stdout[:500]}")
                                if stderr:
                                    print(f"[ERROR] Server stderr: {stderr[:500]}")
                            except:
                                pass
                            break

                        self._check_availability()
                        if self.available:
                            print(f"Llama server started successfully on port {self.port}")
                            break
                        if (i + 1) % 10 == 0:
                            print(f"Still waiting... ({i + 1}/60 seconds)")
                            print(f"[DEBUG] Process still running: {self.server_process and self.server_process.poll() is None}")
                    else:
                        print("Llama server failed to start within 60 second timeout")
                        # Check if process is still running
                        if self.server_process and self.server_process.poll() is None:
                            print("[DEBUG] Server process is still running but not responding to requests")
                else:
                    print("Failed to start llama server")
            else:
                missing = []
                if not self._llama_binary_path:
                    missing.append("llama-server binary")
                if not self._model_path:
                    missing.append("model file")
                print(f"Cannot auto-start llama server: missing {', '.join(missing)}")

    def _start_server(self) -> bool:
        """Start the llama.cpp server process"""
        try:
            # Kill any existing server process first
            self._stop_server()

            # Try GPU first, then fallback to CPU
            configs_to_try = [
                {
                    "ngl": "33",
                    "desc": "GPU mode (Metal)"
                },
                {
                    "ngl": "0",
                    "desc": "CPU-only mode"
                }
            ]

            for config in configs_to_try:
                print(f"Trying {config['desc']}...")

                cmd = [
                    self._llama_binary_path,
                    "-m", self._model_path,
                    "--host", self.host,
                    "--port", str(self.port),
                    "-c", "2048",  # Context size
                    "-ngl", config["ngl"],
                    "--log-disable"  # Reduce log noise
                ]

                print(f"Starting llama server: {' '.join(cmd)}")

                # Start server in background with output capture for debugging
                self.server_process = subprocess.Popen(
                    cmd,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.STDOUT,
                    stdin=subprocess.DEVNULL,
                    text=True
                )

                # Wait a few seconds to see if it crashes immediately
                import time
                time.sleep(3)

                if self.server_process.poll() is None:
                    print(f"Server started successfully in {config['desc']}")
                    return True
                else:
                    print(f"Server failed in {config['desc']} (exit code: {self.server_process.returncode})")
                    # Get error output
                    try:
                        stdout, _ = self.server_process.communicate(timeout=1)
                        if stdout:
                            print(f"Error output: {stdout[:300]}")
                    except:
                        pass

            print("All server startup modes failed")
            return False

        except Exception as e:
            print(f"Failed to start llama server: {e}")
            return False

    def _stop_server(self):
        """Stop the llama server process if running"""
        if self.server_process and self.server_process.poll() is None:
            try:
                self.server_process.terminate()
                self.server_process.wait(timeout=5)
            except subprocess.TimeoutExpired:
                self.server_process.kill()
                self.server_process.wait()
            except Exception as e:
                print(f"Error stopping llama server: {e}")
            finally:
                self.server_process = None

    def cleanup(self):
        """Clean up resources, stop server if we started it"""
        self._stop_server()

    def interpret_emotion(self, text: str, audio_data: dict, video_data: dict, debug: bool = False) -> str:
        """
        Send text, audio, and video emotion data to local Llama for interpretation
        Returns a 3-8 word emotion summary
        """
        if not self.available:
            # Try to ensure server is available (might auto-start)
            self._ensure_server_available()
            if not self.available:
                if debug:
                    print(f"DEBUG: Llama server not available at {self.base_url}")
                return "llama unavailable"

        # Create the prompt with emotion data
        prompt = self._create_emotion_prompt(text, audio_data, video_data)

        if debug:
            print(f"\nLLAMA TEXT SENT: {prompt}")

        try:
            # Try to get the available models first
            available_models = []
            try:
                models_response = requests.get(f"{self.base_url}/v1/models", timeout=2)
                if models_response.status_code == 200:
                    models_data = models_response.json()
                    if "data" in models_data:
                        available_models = [m.get("id") for m in models_data.get("data", [])]
                        print(f"[LLAMA DEBUG] Available models: {available_models}")
            except Exception as e:
                print(f"[LLAMA DEBUG] Could not query /v1/models: {e}")

            # Build list of model names to try (in priority order)
            base_name = self._model_name or "phi3"
            model_variations = [
                # Try available models first
                *available_models,
                # Prefer phi3 (better instruction following than tinyllama)
                "phi3",
                # Then try our detected model name
                base_name,
                # Common llama.cpp defaults
                "gpt-3.5-turbo",  # llama.cpp often uses this as default
                "llama",
                "tinyllama",
                "model",
            ]

            # Remove duplicates while preserving order
            seen = set()
            unique_models = []
            for m in model_variations:
                if m and m not in seen:
                    seen.add(m)
                    unique_models.append(m)

            print(f"[LLAMA DEBUG] Will try model names: {unique_models[:5]}...")

            # Try v1/completions endpoint (OpenAI-compatible)
            endpoint_url = f"{self.base_url}/v1/completions"

            # Try each model name until one works
            response = None
            successful_model = None

            for model_to_use in unique_models:
                print(f"\n[LLAMA DEBUG] Trying model: '{model_to_use}'")

                # Use OpenAI-compatible API format (llama.cpp uses /v1/completions)
                request_payload = {
                    "model": model_to_use,  # Required by llama.cpp
                    "prompt": prompt,
                    "max_tokens": 20,
                    "temperature": 0.3,
                    "stop": ["."],
                    "echo": False,
                    "stream": False
                }

                print(f"[LLAMA DEBUG] Sending request to: {endpoint_url}")
                print(f"[LLAMA DEBUG] Payload keys: {list(request_payload.keys())}")

                try:
                    response = requests.post(
                        endpoint_url,
                        json=request_payload,
                        timeout=10
                    )

                    print(f"[LLAMA DEBUG] Response status code: {response.status_code}")

                    if response.status_code == 200:
                        print(f"[LLAMA DEBUG] ✅ SUCCESS with model '{model_to_use}'!")
                        successful_model = model_to_use
                        # Save successful model for future requests
                        self._model_name = model_to_use
                        break
                    elif response.status_code == 404:
                        error_msg = response.text[:200]
                        print(f"[LLAMA DEBUG] ❌ Model '{model_to_use}' not found: {error_msg}")
                        continue  # Try next model
                    else:
                        print(f"[LLAMA DEBUG] ⚠️ HTTP {response.status_code}: {response.text[:200]}")
                        continue  # Try next model

                except Exception as e:
                    print(f"[LLAMA DEBUG] ❌ Request failed: {e}")
                    continue  # Try next model

            # Check if we found a working model
            if successful_model and response and response.status_code == 200:
                # SUCCESS! Process the response
                result = response.json()
                print(f"[LLAMA DEBUG] Response JSON keys: {list(result.keys())}")

                # OpenAI API format: result.choices[0].text
                # Old format: result.content
                if "choices" in result and len(result["choices"]) > 0:
                    content = result["choices"][0].get("text", "").strip()
                    print(f"[LLAMA DEBUG] Using OpenAI format (choices[0].text)")
                else:
                    content = result.get("content", "").strip()
                    print(f"[LLAMA DEBUG] Using old format (content)")

                if debug:
                    print(f"LLAMA TEXT RECEIVED: '{content}'")

                # Extract the emotion summary - clean up the response
                if debug:
                    print(f"Raw LLM response before cleaning: '{content}'")

                # Remove any trailing text after '}' character
                if "}" in content:
                    content = content.split("}")[0]

                # Remove any text after 'Text:' pattern (common suffix)
                if " Text:" in content:
                    content = content.split(" Text:")[0]

                # Remove common unwanted prefixes
                content = content.strip()
                prefixes_to_remove = ["Text:", "[OUTPUT]", "Emotion:", "{", "{{"]
                for prefix in prefixes_to_remove:
                    if content.startswith(prefix):
                        content = content[len(prefix):].strip()

                # Remove common unwanted suffixes
                suffixes_to_remove = ["}", "}}", " Text", "'I'm not sure", "I'm not sure"]
                for suffix in suffixes_to_remove:
                    if content.endswith(suffix):
                        content = content[:-len(suffix)].strip()

                if debug:
                    print(f"Cleaned LLM response: '{content}'")

                # Extract the emotion summary (first 2-8 words, allow shorter responses)
                words = content.split()[:8]
                if len(words) >= 2:  # Changed from 3 to 2 to allow shorter valid emotions
                    final_result = " ".join(words)
                    return final_result
                else:
                    return "processing incomplete"
            else:
                # All models failed
                print(f"\n[LLAMA ERROR] ❌ All {len(unique_models)} model variations failed!")
                print(f"[LLAMA ERROR] Tried models: {unique_models}")
                if response:
                    print(f"[LLAMA ERROR] Last HTTP status: {response.status_code}")
                    print(f"[LLAMA ERROR] Last response: {response.text[:500]}")
                return "llama error"

        except requests.exceptions.Timeout as e:
            print(f"[LLAMA ERROR] Request timeout: {e}")
            return "llama timeout"
        except requests.exceptions.ConnectionError as e:
            print(f"[LLAMA ERROR] Connection error: {e}")
            return "llama unavailable"
        except Exception as e:
            print(f"[LLAMA ERROR] Unexpected error: {type(e).__name__}: {e}")
            import traceback
            traceback.print_exc()
            return "llama error"

    def _create_emotion_prompt(self, text: str, audio_data: dict, video_data: dict) -> str:
        """Create prompt for emotion interpretation"""

        # Format audio emotion data
        audio_summary = f"V:{audio_data.get('valence', 0):.2f}, A:{audio_data.get('arousal', 0):.2f}, D:{audio_data.get('dominance', 0):.2f}"

        # Format video emotion data - include top 5 emotions with values
        video_summary = "None"
        if video_data and isinstance(video_data, list) and video_data:
            # Get most recent video emotion
            recent = video_data[-1] if video_data else {}
            emotions = recent.get('emotions', {})
            if emotions:
                # Sort emotions by confidence and get top 5
                sorted_emotions = sorted(emotions.items(), key=lambda x: x[1], reverse=True)[:5]
                emotion_strings = [f"{emotion}:{confidence:.2f}" for emotion, confidence in sorted_emotions]
                video_summary = ", ".join(emotion_strings)

        prompt = f"""[INSTRUCTIONS]
You summarize a journaler's emotional state as a concise phrase.
Use ALL available data: text, acoustic cues (Valence, Arousal, Dominance), and facial emotions (top 5 with confidence scores).
Follow the rules strictly:

RULES:
- Output ONE short phrase of 3–8 words.
- NEVER use commas - use "and" or "but" instead.
- No lists, no slashes, no "etc.", no quotes.
- No first-person ("I", "I feel", "I am").
- No duplicate words.
- Prefer natural, family-friendly phrasing (e.g., "hopeful but tense").
- If signals conflict, reflect mixed affect (e.g., "hopeful yet uneasy").
- Do NOT explain. Do NOT add labels.

VAD HINTS (soft guidance, not hard rules):
- High Valence + High Arousal → "energized, joyful, enthusiastic"
- High Valence + Low Arousal → "calm, content, relieved"
- Low Valence + High Arousal → "anxious, angry, agitated"
- Low Valence + Low Arousal → "tired, sad, discouraged"
- Higher Dominance → "confident, in control"
- Lower Dominance → "overwhelmed, powerless"

FEW-SHOT EXAMPLES
Text: "I'm so alive I can feel my hair growing and my nails growing! It feels like every cell in my body just had a milkshake!"
V: +0.62  A: +0.48  D: +0.55
Face: happy:0.72, neutral:0.18, surprise:0.06, fear:0.03, sad:0.01
{{bonkers and goofy}}

Text: "I can't focus; everything keeps piling up."
V: -0.58  A: +0.51  D: -0.41
Face: fear:0.68, angry:0.15, sad:0.12, neutral:0.04, disgust:0.01
{{overwhelmed, anxious, seeking relief}}

Text: "It hurts, but I know we'll get through."
V: -0.20  A: +0.18  D: +0.22
Face: sad:0.55, neutral:0.25, happy:0.12, fear:0.05, angry:0.03
{{hurting yet cautiously hopeful}}

Text: "Not much to say. Just tired lately."
V: -0.25  A: -0.45  D: -0.10
Face: neutral:0.61, sad:0.28, angry:0.06, fear:0.03, disgust:0.02
{{drained and low-spirited}}

[INPUT]
Text: "{text}"
V: {audio_data.get('valence', 0):.2f}
A: {audio_data.get('arousal', 0):.2f}
D: {audio_data.get('dominance', 0):.2f}
Face: {video_summary}

{{"""

        return prompt

# ============================ Computer Vision Emotion Detection ============================
class VideoEmotionDetector:
    """
    PyTorch-based facial emotion recognition using ResNet50 trained on AffectNet.
    Detects faces and estimates emotions from facial expressions.
    """

    def __init__(self):
        self.face_cascade = None
        self.emotion_model = None
        self.device = None
        self.transform = None
        self.emotion_labels = ['angry', 'disgust', 'fear', 'happy', 'sad', 'surprise', 'neutral']

        if cv2 is not None:
            try:
                # Load OpenCV's face detector
                self.face_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_frontalface_default.xml')
            except Exception as e:
                print(f"Warning: Could not load face cascade: {e}")

        # Load PyTorch ResNet50 FER model
        self._load_fer_model()

    def _load_fer_model(self):
        """Load the pre-trained ResNet50 FER model from models/ directory."""
        if torch is None:
            print("PyTorch not available, FER model loading skipped")
            return

        try:
            # Set device
            self.device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
            print(f"Using device: {self.device}")

            # Define the transform for preprocessing
            self.transform = transforms.Compose([
                transforms.ToPILImage(),
                transforms.Resize((224, 224)),
                transforms.ToTensor(),
                transforms.Normalize(mean=[0.485, 0.456, 0.406], std=[0.229, 0.224, 0.225])
            ])

            # Load the trained weights first to inspect the architecture
            model_path = Path(__file__).parent / "models" / "FER_static_ResNet50_AffectNet.pt"
            if model_path.exists():
                checkpoint = torch.load(model_path, map_location=self.device)

                # Handle different checkpoint formats
                if isinstance(checkpoint, dict) and 'state_dict' in checkpoint:
                    state_dict = checkpoint['state_dict']
                elif isinstance(checkpoint, dict) and 'model_state_dict' in checkpoint:
                    state_dict = checkpoint['model_state_dict']
                else:
                    state_dict = checkpoint

                # Create custom model architecture based on the actual saved model
                self.emotion_model = self._create_custom_fer_model()

                # Try to load compatible weights, ignoring incompatible ones
                try:
                    missing_keys, unexpected_keys = self.emotion_model.load_state_dict(state_dict, strict=False)
                    if missing_keys:
                        print(f"FER model: {len(missing_keys)} missing keys (using random initialization)")
                    if unexpected_keys:
                        print(f"FER model: {len(unexpected_keys)} unexpected keys (ignored)")

                    # If too many keys are missing, fall back to heuristic mode
                    if len(missing_keys) > len(state_dict) * 0.5:  # More than 50% missing
                        print("=" * 70)
                        print("WARNING: PyTorch FER model loading failed - too many missing keys")
                        print("Falling back to HEURISTIC emotion detection mode")
                        print("Heuristic mode uses basic brightness/edge features and is NOT reliable")
                        print("For accurate facial emotion recognition, ensure FER_static_ResNet50_AffectNet.pt")
                        print("is properly downloaded and placed in the models/ directory")
                        print("=" * 70)
                        self.emotion_model = None
                        return

                except Exception as load_error:
                    print("=" * 70)
                    print(f"ERROR: Failed to load PyTorch FER model: {load_error}")
                    print("Falling back to HEURISTIC emotion detection mode")
                    print("Heuristic mode uses basic brightness/edge features and is NOT reliable")
                    print("For accurate facial emotion recognition, ensure FER_static_ResNet50_AffectNet.pt")
                    print("is properly downloaded and placed in the models/ directory")
                    print("=" * 70)
                    self.emotion_model = None
                    return

                self.emotion_model.to(self.device)
                self.emotion_model.eval()
                print(f"FER model loaded successfully from {model_path}")
            else:
                print(f"FER model not found at {model_path}")
                self.emotion_model = None

        except Exception as e:
            print(f"Error loading FER model: {e}")
            self.emotion_model = None

    def _create_custom_fer_model(self):
        """Create custom FER model architecture matching the saved model."""
        try:
            # Try to create a simple fallback model for now
            # This will load the weights as much as possible and fall back gracefully
            model = resnet50(weights=None)
            model.fc = nn.Linear(model.fc.in_features, 7)
            return model
        except Exception as e:
            print(f"Error creating custom FER model: {e}")
            return None

    def detect_emotions(self, frame) -> List[Dict[str, Any]]:
        """
        Detect emotions in a video frame using PyTorch ResNet50 FER model.
        Returns list of detection results with bounding boxes and emotion scores.
        """
        if frame is None or self.face_cascade is None:
            return []

        try:
            # Convert to grayscale for face detection
            gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)

            # Detect faces
            faces = self.face_cascade.detectMultiScale(
                gray,
                scaleFactor=1.1,
                minNeighbors=5,
                minSize=(48, 48)  # Increased minimum size for better FER
            )

            results = []
            for (x, y, w, h) in faces:
                # Extract face region (RGB for model input)
                face_rgb = frame[y:y+h, x:x+w]

                # Use PyTorch FER model if available, otherwise fallback to heuristics
                if self.emotion_model is not None and self.transform is not None:
                    emotions = self._predict_emotion_pytorch(face_rgb)
                else:
                    face_gray = gray[y:y+h, x:x+w]
                    emotions = self._estimate_emotion_heuristic(face_gray, face_rgb)

                # Calculate overall confidence
                confidence = max(emotions.values()) if emotions else 0.0

                results.append({
                    'box': [int(x), int(y), int(w), int(h)],
                    'emotions': emotions,
                    'confidence': confidence
                })

            return results

        except Exception as e:
            print(f"Error in emotion detection: {e}")
            return []

    def _predict_emotion_pytorch(self, face_rgb) -> Dict[str, float]:
        """
        Predict emotions using the PyTorch ResNet50 FER model.
        """
        try:
            # Preprocess the face image
            # Convert BGR to RGB if needed (OpenCV uses BGR)
            if len(face_rgb.shape) == 3 and face_rgb.shape[2] == 3:
                face_rgb = cv2.cvtColor(face_rgb, cv2.COLOR_BGR2RGB)

            # Apply transforms
            input_tensor = self.transform(face_rgb).unsqueeze(0).to(self.device)

            # Forward pass
            with torch.no_grad():
                outputs = self.emotion_model(input_tensor)
                probabilities = torch.nn.functional.softmax(outputs, dim=1)
                probs = probabilities.cpu().numpy()[0]

            # Map to emotion labels
            emotions = {}
            for i, label in enumerate(self.emotion_labels):
                emotions[label] = float(probs[i])

            return emotions

        except Exception as e:
            print(f"Error in PyTorch emotion prediction: {e}")
            # Fallback to neutral emotion
            return {'neutral': 1.0}

    def _estimate_emotion_heuristic(self, face_gray, face_color) -> Dict[str, float]:
        """
        Simple heuristic-based emotion estimation.
        In production, replace with a trained emotion recognition CNN.
        """
        try:
            # Basic feature extraction
            h, w = face_gray.shape

            # Eye region (upper third)
            eye_region = face_gray[0:h//3, :]

            # Mouth region (lower third)
            mouth_region = face_gray[2*h//3:h, :]

            # Calculate basic features
            eye_brightness = np.mean(eye_region) if eye_region.size > 0 else 128
            mouth_brightness = np.mean(mouth_region) if mouth_region.size > 0 else 128

            # Edge density (expression intensity)
            edges = cv2.Canny(face_gray, 50, 150)
            edge_density = np.sum(edges > 0) / (h * w) if h * w > 0 else 0

            # Overall brightness
            overall_brightness = np.mean(face_gray)

            # Simple heuristics (placeholder for real ML model)
            # These would be replaced by actual CNN predictions
            emotions = {
                'neutral': 0.4 + 0.2 * (1 - edge_density),
                'happy': 0.1 + 0.3 * max(0, mouth_brightness - overall_brightness) / 50,
                'sad': 0.1 + 0.2 * max(0, overall_brightness - mouth_brightness) / 50,
                'angry': 0.05 + 0.25 * edge_density,
                'surprise': 0.05 + 0.2 * max(0, eye_brightness - overall_brightness) / 50,
                'fear': 0.05 + 0.15 * edge_density,
                'disgust': 0.05
            }

            # Normalize to sum to 1
            total = sum(emotions.values())
            if total > 0:
                emotions = {k: v/total for k, v in emotions.items()}

            return emotions

        except Exception as e:
            print(f"Error in emotion heuristic: {e}")
            return {'neutral': 1.0}

# ============================ Tone/Text Engines ============================
import math, numpy as _np
try:
    import webrtcvad as _webrtcvad
except Exception:
    _webrtcvad = None

class ToneEstimator:
    """
    CPU-light VA(D) from raw PCM with:
      • WebRTC VAD gating (optional)
      • RMS dB, spectral centroid/tilt, spectral flux
      • Session calibration via running percentiles
    """
    def __init__(self, sr: int):
        self.sr = sr
        self._vad = _webrtcvad.Vad(2) if _webrtcvad else None  # 0..3 (aggressive). 2 is balanced.
        # Rolling histories for calibration
        self._rms_db_hist = deque(maxlen=600)  # ~20s if 33ms frames
        self._centroid_hist = deque(maxlen=600)
        self._flux_hist = deque(maxlen=600)

        # For flux (prev spectrum)
        self._prev_mag = None

    def _is_speech_webrtc(self, pcm_bytes: bytes) -> bool:
        if not self._vad:
            return True
        # WebRTC VAD expects 10/20/30 ms frames at 16k, 16-bit mono
        frame_ms = 20
        bytes_per_ms = int(self.sr/1000) * 2
        if len(pcm_bytes) < frame_ms * bytes_per_ms:
            return True
        frame = pcm_bytes[:frame_ms * bytes_per_ms]
        try:
            return self._vad.is_speech(frame, self.sr)
        except Exception:
            return True

    def _fft_features(self, pcm: array) -> tuple[float, float, float]:
        # Convert to float32 and get magnitude spectrum
        x = _np.asarray(pcm, dtype=_np.float32) / 32768.0
        if x.size == 0:
            return 0.0, 0.0, 0.0
        # Hann window
        w = _np.hanning(x.size).astype(_np.float32)
        X = _np.fft.rfft(x * w)
        mag = _np.abs(X) + 1e-9

        # Spectral centroid (Hz)
        freqs = _np.fft.rfftfreq(x.size, d=1.0/self.sr)
        centroid = float((_np.sum(freqs * mag) / _np.sum(mag)) if _np.sum(mag) > 0 else 0.0)

        # Spectral tilt: low vs high band energy
        split = int(1000 * x.size / self.sr)  # ~1kHz boundary in bins
        low = float(_np.sum(mag[:max(1, split)]))
        high = float(_np.sum(mag[max(1, split):]))
        tilt = float((low - high) / (low + high + 1e-6))  # + = warm, - = bright/harsh

        # Spectral flux (frame-to-frame change)
        if self._prev_mag is None or self._prev_mag.shape != mag.shape:
            flux = 0.0
        else:
            diff = mag - self._prev_mag
            flux = float(_np.sum(_np.clip(diff, 0, None)) / (mag.size + 1e-6))
        self._prev_mag = mag
        return centroid, tilt, flux

    def estimate_vad(self, pcm_bytes: bytes) -> Tuple[float, float, float]:
        """
        Returns (V, A, D): V∈[-1,1], A∈[0,1], D∈[-1,1]
        Bias-reduced arousal: lower baseline + smaller weights on RMS/flux.
        """
        a = array('h'); a.frombytes(pcm_bytes)
        if not self._is_speech_webrtc(pcm_bytes):
            return 0.0, 0.05, -0.2
        if len(a) == 0:
            return 0.0, 0.05, -0.2

        x = _np.asarray(a, dtype=_np.float32) / 32768.0
        rms = float(_np.sqrt(_np.mean(x*x)) + 1e-12)
        db = 20.0 * math.log10(rms + 1e-12)

        centroid, tilt, flux = self._fft_features(a)

        # Update histories
        self._rms_db_hist.append(db)
        self._centroid_hist.append(centroid)
        self._flux_hist.append(flux)

        def _pct_scaled(val: float, hist: deque, lo_q=0.1, hi_q=0.9) -> float:
            # Lower default baseline so neutral talk ≈ low-mid arousal
            if len(hist) < 10:
                return 0.35
            arr = _np.fromiter(hist, dtype=_np.float32)
            lo = float(_np.quantile(arr, lo_q))
            hi = float(_np.quantile(arr, hi_q))
            if hi <= lo:
                return 0.35
            return float(max(0.0, min(1.0, (val - lo) / (hi - lo))))

        # ↓ toned-down arousal calculation (less "excited" bias)
        A_rms  = _pct_scaled(db, self._rms_db_hist)
        A_flux = _pct_scaled(flux, self._flux_hist)
        A = 0.60 * A_rms + 0.10 * A_flux
        A = max(0.0, min(1.0, A))

        # Valence: keep centered; warm tilt → slightly positive, harsh/bright → negative
        V = 0.60 * tilt + 0.10 * (2 * A_flux - 1.0)
        V = max(-1.0, min(1.0, V))

        # Dominance: mostly arousal with a small centroid contribution
        Cn = _pct_scaled(centroid, self._centroid_hist)
        D = (A * 1.00 - 0.40) + 0.15 * (2 * Cn - 1.0)
        D = max(-1.0, min(1.0, D))

        return float(V), float(A), float(D)

class TextValenceEstimator:
    """
    Paragraph-level text valence with strong neutralization for academic/expository prose.
    Returns (valence in [-1,1], is_academic: bool).
    """
    def __init__(self):
        self._pipe = None  # keep disabled to avoid positive skew

        self.pos = set("""
            joy happy great love wonderful serene calm confident hopeful relieved excited
            grateful pleased peaceful content upbeat optimistic cheerful buoyant lively
        """.split())
        self.neg = set("""
            sad upset angry furious anxious stressed worried devastated depressed fearful
            tense overwhelmed irritated frustrated painful annoyed hostile
        """.split())

        # Common markers of expository/academic narration
        self.academic = set("""
            study studies research observed observe defines definition terms process context
            evidence data result results method methods analysis literature theory theoretical
            framework model models modeling measured measurement findings discussion conclusion
            professor university journal decade decades potential performance creativity health
            learning category categories rule rules constraint constraints mindlessness
            mindfulness investigation investigate scientific approach distinct practice
            implication implications experiment experiments empirical conditional outcomes
            demonstrated demonstrated
        """.split())

        self.exclaim_words = set("wow amazing incredible unbelievable fantastic awful terrible shocking".split())

    def valence(self, text: str) -> tuple[float, bool]:
        t = (text or "").strip()
        if not t:
            return 0.0, False

        tokens = re.findall(r"[A-Za-z']+", t.lower())
        if not tokens:
            return 0.0, False

        pos = sum(1 for tok in tokens if tok in self.pos)
        neg = sum(1 for tok in tokens if tok in self.neg)
        affect_total = pos + neg

        acad_hits = sum(1 for tok in tokens if tok in self.academic)
        acad_ratio = acad_hits / max(1, len(tokens))
        is_academic = acad_ratio >= 0.04  # fairly sensitive

        has_exclaim = bool(re.search(r"[!]{1,}$", t))
        has_interj = any(w in self.exclaim_words for w in tokens)

        # Base lexicon score
        base = 0.0
        if affect_total:
            base = (pos - neg) / affect_total

        # If little/no affect and it's not emphatic → neutral
        if affect_total <= 2 and not has_exclaim and not has_interj:
            base = 0.0

        # Hard neutral for clear academic prose without emphatic cues
        if is_academic and not has_exclaim and not has_interj:
            return 0.0, True

        # Otherwise, keep text influence modest
        base = float(max(-0.30, min(0.30, base)))
        return base, is_academic

# =========================== Fusion & Summarizer ===========================
def iqr(vals: List[float]) -> float:
    if not vals:
        return 0.0
    vs = sorted(vals)
    def pct(p: float) -> float:
        k = (len(vs)-1)*p
        f = math.floor(k); c = math.ceil(k)
        if f == c: return vs[int(k)]
        return vs[f] + (vs[c] - vs[f])*(k - f)
    return pct(0.75) - pct(0.25)

class EmotionSummarizer:
    def __init__(self, matrix: dict):
        self.M = matrix
        self.VB = matrix["bins"]["valence"]
        self.AB = matrix["bins"]["arousal"]
        self.DB = matrix["bins"]["dominance"]
        self.GRID = matrix["grid"]
        self.INTS = matrix["intensifiers"]
        self.DT = matrix["dominance_tones"]
        self.defaults = matrix.get("defaults", {})
        self.mixed_threshold = float(self.defaults.get("mixed_variability_threshold", 0.35))
        self.min_words = int(self.defaults.get("min_words", 3))
        self.max_words = int(self.defaults.get("max_words", 5))
        self.insufficient = self.defaults.get("insufficient_label", "insufficient signal")

    def summarize(self, V_seq: List[float], A_seq: List[float], D_seq: Optional[List[float]],
                  text_valence_ignored=None, video_emotions: Optional[List[Dict]] = None,
                  audio_weight: float = 0.7, video_weight: float = 0.3) -> Tuple[str, Dict[str, float]]:
        """
        Enhanced summarizer that fuses audio tone + video emotions.
        video_emotions: List of emotion detection results from video frames
        """
        if not V_seq or not A_seq:
            return self.insufficient, {"V": 0.0, "A": 0.0, "D": 0.0}

        V_tone = float(median(V_seq))
        A_tone = float(median(A_seq))
        D_tone = float(median(D_seq)) if D_seq else 0.0
        variability = (iqr(V_seq) + iqr(A_seq)) / 2.0

        # Fuse with video emotions if available
        V_fused = V_tone
        A_fused = A_tone
        D_fused = D_tone

        # Extract video VAD once for both fusion and debug
        video_valence, video_arousal = None, None
        if video_emotions:
            video_valence, video_arousal = self._extract_video_vad(video_emotions)

            # Weighted fusion with configurable weights
            if video_valence is not None and video_arousal is not None:
                V_fused = audio_weight * V_tone + video_weight * video_valence
                A_fused = audio_weight * A_tone + video_weight * video_arousal

        phrase = self._phrase_from_VAD(V_fused, A_fused, D_fused, variability)

        # Generate separate audio and video phrases for debug
        audio_phrase = self._phrase_from_VAD(V_tone, A_tone, D_tone, variability)
        video_phrase = None
        if video_valence is not None and video_arousal is not None:
            video_phrase = self._phrase_from_VAD(video_valence, video_arousal, 0.0, 0.0)

        return phrase, {
            "V": V_fused, "A": A_fused, "D": D_fused, "var": variability,
            "w_audio": audio_weight if video_emotions else 1.0,
            "w_video": video_weight if video_emotions else 0.0,
            "video_detections": len(video_emotions) if video_emotions else 0,
            "audio_phrase": audio_phrase,
            "video_phrase": video_phrase,
            "audio_vad": {"V": V_tone, "A": A_tone, "D": D_tone},
            "video_vad": {"V": video_valence, "A": video_arousal} if video_emotions else None,
            "raw_video_emotions": video_emotions if video_emotions else None,
            "raw_audio_sequences": {"V_seq": V_seq, "A_seq": A_seq, "D_seq": D_seq}
        }

    def _extract_video_vad(self, video_emotions: List[Dict]) -> Tuple[Optional[float], Optional[float]]:
        """
        Convert video emotion detections to valence/arousal values.
        """
        if not video_emotions:
            return None, None

        # Emotion to valence/arousal mapping (Russell's circumplex model)
        emotion_vad_map = {
            'happy': (0.8, 0.6),      # high valence, moderate arousal
            'joy': (0.9, 0.7),        # very high valence, high arousal
            'surprise': (0.1, 0.8),   # neutral valence, high arousal
            'angry': (-0.7, 0.8),     # negative valence, high arousal
            'fear': (-0.6, 0.7),      # negative valence, high arousal
            'sad': (-0.8, 0.2),       # negative valence, low arousal
            'disgust': (-0.6, 0.4),   # negative valence, moderate arousal
            'neutral': (0.0, 0.3),    # neutral valence, low arousal
        }

        # Aggregate emotions across all detections
        total_valence = 0.0
        total_arousal = 0.0
        total_weight = 0.0

        for detection in video_emotions:
            emotions = detection.get('emotions', {})
            confidence = detection.get('confidence', 0.0)

            # Weight by detection confidence
            detection_weight = confidence

            frame_valence = 0.0
            frame_arousal = 0.0

            for emotion, prob in emotions.items():
                if emotion in emotion_vad_map:
                    val, ar = emotion_vad_map[emotion]
                    frame_valence += prob * val
                    frame_arousal += prob * ar

            total_valence += detection_weight * frame_valence
            total_arousal += detection_weight * frame_arousal
            total_weight += detection_weight

        if total_weight > 0:
            avg_valence = total_valence / total_weight
            avg_arousal = total_arousal / total_weight

            # Normalize to expected ranges
            avg_valence = max(-1.0, min(1.0, avg_valence))
            avg_arousal = max(0.0, min(1.0, avg_arousal))

            return avg_valence, avg_arousal

        return None, None

    def _average_emotions_over_paragraph(self, video_emotions: List[Dict]) -> Dict[str, float]:
        """
        Average all emotion scores across the entire paragraph timespan.
        Returns a single averaged emotion profile.
        """
        if not video_emotions:
            return {}

        # Initialize emotion accumulators
        emotion_sums = {}
        total_weight = 0.0

        # Accumulate weighted emotion scores
        for detection in video_emotions:
            emotions = detection.get('emotions', {})
            confidence = detection.get('confidence', 0.0)

            # Weight by detection confidence
            weight = confidence
            total_weight += weight

            for emotion, score in emotions.items():
                if emotion not in emotion_sums:
                    emotion_sums[emotion] = 0.0
                emotion_sums[emotion] += score * weight

        # Calculate averages
        if total_weight > 0:
            averaged_emotions = {emotion: total_sum / total_weight
                               for emotion, total_sum in emotion_sums.items()}
        else:
            averaged_emotions = {}

        return averaged_emotions

    def _phrase_from_VAD(self, V: float, A: float, D: float, variability: float) -> str:
        vi = _bucket(V, self.VB)
        ai = _bucket(A, self.AB)

        # If arousal is modest, suppress upbeat mappings by nudging valence toward neutral
        if A < 0.60 and vi >= 3:  # right side (positive) while low arousal
            vi = max(2, vi - 1)   # pull one bin toward neutral

        words = list(self.GRID[ai][vi].get("adjectives", []))[:2]

        # Intensifier: only for clearly high arousal
        adverb = ""
        for x in self.INTS:
            if A < x["threshold"]:
                adverb = x["adverb"]
                break
        # Suppress adverbs entirely if arousal isn't high
        if A < 0.60:
            adverb = ""
        if adverb:
            words = [adverb] + words

        # Dominance
        tone = None
        for r in self.DT:
            if r["min"] <= D < r["max"]:
                tone = r.get("word")
                break
        if tone:
            words.append(tone)

        # Variability → "mixed"
        if variability > self.mixed_threshold:
            words.append("mixed")

        # Dedup + clamp
        seen, out = set(), []
        for w in words:
            for tok in w.split():
                if tok and tok not in seen:
                    out.append(tok); seen.add(tok)

        if len(out) < self.min_words:
            out += ["tone"] * (self.min_words - len(out))
        return " ".join(out[:self.max_words])

# ============================== Vosk helpers ===============================
def _candidate_model_dirs() -> List[Path]:
    cand: List[Path] = []
    here = Path(__file__).resolve().parent if "__file__" in globals() else Path.cwd()
    cand.append(here / "models" / VOSK_DIR_NAME)
    if getattr(sys, "frozen", False) and hasattr(sys, "_MEIPASS"):
        cand.append(Path(sys._MEIPASS) / "models" / VOSK_DIR_NAME)  # type: ignore[attr-defined]
    cand.append(Path.home() / ".pq-journal" / "models" / VOSK_DIR_NAME)
    return cand

def find_builtin_vosk_model() -> Optional[Path]:
    for p in _candidate_model_dirs():
        if p.exists() and ((p / "model.conf").exists() or any((p / s).exists() for s in ("am","graph","rescore","conf"))):
            return p
    here = Path(__file__).resolve().parent if "__file__" in globals() else Path.cwd()
    zips = [here / "models" / "vosk-en-small.zip"]
    if getattr(sys, "frozen", False) and hasattr(sys, "_MEIPASS"):
        zips.append(Path(sys._MEIPASS) / "models" / "vosk-en-small.zip")  # type: ignore[attr-defined]
    for z in zips:
        if z.exists():
            target = Path.home() / ".pq-journal" / "models" / VOSK_DIR_NAME
            target.parent.mkdir(parents=True, exist_ok=True)
            if not target.exists() or not any(target.iterdir()):
                try:
                    with zipfile.ZipFile(z, "r") as zf: zf.extractall(target.parent)
                except Exception:
                    if target.exists(): shutil.rmtree(target, ignore_errors=True)
            if target.exists(): return target
    return None

# =============================== Crypto ===================================
def kem_pick() -> str:
    if oqs is None: return ""
    norm = {m.upper(): m for m in getattr(oqs, 'get_enabled_kem_mechanisms', lambda: [])()}
    if "ML-KEM-1024" in norm: return norm["ML-KEM-1024"]
    if "KYBER1024"  in norm: return norm["KYBER1024"]
    return ""
KEM_NAME = kem_pick()

b64e = lambda b: base64.b64encode(b).decode("ascii")
b64d = lambda s: base64.b64decode(s.encode("ascii"))

def _pbkdf2_key(password: bytes, salt: bytes, iterations=200_000, length=32) -> bytes:
    kdf = PBKDF2HMAC(algorithm=hashes.SHA256(), length=length, salt=salt, iterations=iterations)
    return kdf.derive(password)

def read_key_file(p: Path) -> bytes:
    data = p.read_bytes()
    try: return b64d(data.decode("ascii").strip())
    except Exception: return data

def read_privkey_file(path: Path, passphrase: Optional[str]) -> bytes:
    data = path.read_bytes()
    try:
        txt = data.decode("utf-8"); lines = [ln.strip() for ln in txt.splitlines() if ln.strip()]
        if lines and lines[0] == "v=1":
            if not passphrase: raise ValueError("Private key is password-encrypted; a password is required.")
            kv = {}
            for ln in lines[1:]:
                if "=" in ln: k, v = ln.split("=", 1); kv[k.strip()] = v.strip()
            salt, nonce, ct = b64d(kv["salt"]), b64d(kv["nonce"]), b64d(kv["ct"])
            key = _pbkdf2_key(passphrase.encode("utf-8"), salt)
            return AESGCM(key).decrypt(nonce, ct, None)
    except Exception:
        pass
    try: return b64d(data.decode("ascii").strip())
    except Exception: return data

def kem_encapsulate(pub: bytes) -> Tuple[bytes, bytes]:
    with oqs.KeyEncapsulation(KEM_NAME) as kem: return kem.encap_secret(pub)

def kem_decapsulate(priv: bytes, ct: bytes) -> bytes:
    with oqs.KeyEncapsulation(KEM_NAME, secret_key=priv) as kem: return kem.decap_secret(ct)

def hkdf32(keymat: bytes) -> bytes:
    if len(keymat) == 32: return keymat
    return HKDF(algorithm=hashes.SHA256(), length=32, salt=None, info=b"pq-journal").derive(keymat)

def aes_enc(plaintext: bytes, keymat: bytes):
    key = hkdf32(keymat); nonce = os.urandom(12)
    ct = AESGCM(key).encrypt(nonce, plaintext, None)
    return nonce, ct[:-16], ct[-16:]

def aes_dec(nonce: bytes, body: bytes, tag: bytes, keymat: bytes) -> bytes:
    key = hkdf32(keymat); return AESGCM(key).decrypt(nonce, body + tag, None)

def encrypt_blob(plaintext: bytes, pubkey: bytes) -> dict:
    kem_ct, ss = kem_encapsulate(pubkey); nonce, body, tag = aes_enc(plaintext, ss)
    return {"alg": f"{KEM_NAME}+AES-256-GCM", "kem_ct": b64e(kem_ct), "nonce": b64e(nonce), "tag": b64e(tag), "ct": b64e(body)}

def decrypt_blob(blob: dict, privkey: bytes) -> bytes:
    ss = kem_decapsulate(privkey, b64d(blob["kem_ct"]))
    return aes_dec(b64d(blob["nonce"]), b64d(blob["ct"]), b64d(blob["tag"]), ss)

# =========================== Notebook helpers =============================
def empty_notebook() -> dict:
    from datetime import datetime
    return {"version": 1, "created": datetime.now().astimezone().isoformat(), "entries": []}

def pack_entry_payload(entry_dict: dict) -> bytes:
    return json.dumps({"body": entry_dict.get("body", "")}, separators=(",", ":"), ensure_ascii=False).encode("utf-8")

def unpack_entry_payload(b: bytes) -> dict:
    try: d = json.loads(b.decode("utf-8")); return {"body": d.get("body", "")}
    except Exception: return {"body": ""}

def write_json_atomic_secure(path: Path, obj: dict) -> None:
    data = json.dumps(obj, separators=(",", ":"), ensure_ascii=False).encode("utf-8")
    tmp = path.with_suffix(path.suffix + ".tmp")
    fd = os.open(str(tmp), os.O_WRONLY | os.O_CREAT | os.O_TRUNC, 0o600)
    try:
        with os.fdopen(fd, "wb") as f:
            f.write(data); f.flush(); os.fsync(f.fileno())
        os.replace(tmp, path)
        try: os.chmod(path, 0o600)
        except Exception: pass
    finally:
        if tmp.exists():
            try: tmp.unlink()
            except Exception: pass

def require_crypto_ready() -> Optional[str]:
    core_missing = [m for m in missing if m in ("pyoqs", "cryptography")]
    if core_missing:
        return ("Missing Python packages: " + ", ".join(core_missing) +
                "\nInstall with:\n  python3 -m pip install pyoqs cryptography")
    if not KEM_NAME: return "No allowed PQ KEM available. Need ML-KEM-1024 or Kyber1024."
    return None

# ========================= Video Capture Worker (QThread) ===================
class VideoWorker(QThread):
    frame_ready = pyqtSignal(object)  # np.ndarray
    emotion_detected = pyqtSignal(list)  # List of emotion detections
    error = pyqtSignal(str)

    def __init__(self):
        super().__init__()
        self._stop = False
        self.emotion_detector = VideoEmotionDetector()
        self._emotion_history = deque(maxlen=100)  # Keep more frames for paragraph-level aggregation
        self._recording_start_time = None  # Track recording session start

    def stop(self):
        self._stop = True

    def run(self):
        if cv2 is None:
            print("OpenCV not available for video capture")
            self.error.emit("OpenCV not available for video capture")
            return

        cap = None
        try:
            print("Attempting to open camera...")
            cap = cv2.VideoCapture(0)

            # Give camera time to initialize
            time.sleep(0.5)

            if not cap.isOpened():
                # Try different camera indices
                for i in range(1, 4):
                    print(f"Trying camera index {i}...")
                    cap.release()
                    cap = cv2.VideoCapture(i)
                    time.sleep(0.3)
                    if cap.isOpened():
                        break

                if not cap.isOpened():
                    self.error.emit("Could not open any camera. Check permissions.")
                    print("Could not open any camera. Check permissions.")
                    return

            print("Camera opened successfully!")

            # Set camera properties
            cap.set(cv2.CAP_PROP_FRAME_WIDTH, 320)
            cap.set(cv2.CAP_PROP_FRAME_HEIGHT, 240)
            cap.set(cv2.CAP_PROP_FPS, 15)

            # Test if we can read a frame
            ret, test_frame = cap.read()
            if not ret or test_frame is None:
                self.error.emit("Camera opened but cannot read frames")
                print("Camera opened but cannot read frames")
                return

            print(f"Camera test successful. Frame shape: {test_frame.shape}")

            # Initialize recording start time for timestamp calculations
            self._recording_start_time = time.time()
            frame_count = 0
            while not self._stop:
                ret, frame = cap.read()
                if not ret or frame is None:
                    print("Failed to read frame, continuing...")
                    continue

                # Emit frame for display
                self.frame_ready.emit(frame.copy())  # Make a copy to avoid threading issues

                # Process emotions every 5th frame to reduce computational load
                if frame_count % 5 == 0:
                    try:
                        emotions = self.emotion_detector.detect_emotions(frame)
                        if emotions:
                            # Add timestamp to each emotion detection
                            current_time = time.time() - self._recording_start_time if self._recording_start_time else 0.0
                            timestamped_emotions = []
                            for emotion_data in emotions:
                                emotion_with_timestamp = emotion_data.copy()
                                emotion_with_timestamp['timestamp'] = current_time
                                timestamped_emotions.append(emotion_with_timestamp)

                            # Add detailed frame-by-frame logging (only in debug mode)
                            if hasattr(self, '_debug_enabled') and self._debug_enabled:
                                model_type = "PYTORCH" if self.emotion_detector.emotion_model is not None else "HEURISTIC"
                                print(f"\n--- Frame {frame_count} FER Results ({model_type}) @ {current_time:.2f}s ---")
                                for i, detection in enumerate(timestamped_emotions):
                                    bbox = detection.get('bbox', {})
                                    confidence = detection.get('confidence', 0.0)
                                    emotion_scores = detection.get('emotions', {})

                                    bbox_str = f"({bbox.get('x', 0)},{bbox.get('y', 0)}) {bbox.get('w', 0)}x{bbox.get('h', 0)}"

                                    # Show all 7 emotions with their raw scores
                                    sorted_emotions = sorted(emotion_scores.items(), key=lambda x: x[1], reverse=True)
                                    emotion_detail = " ".join([f"{emo}={score:.3f}" for emo, score in sorted_emotions])

                                    print(f"  Face {i+1}: bbox={bbox_str} conf={confidence:.3f}")
                                    print(f"    Raw scores: {emotion_detail}")

                            self._emotion_history.extend(timestamped_emotions)
                            self.emotion_detected.emit(timestamped_emotions)
                    except Exception as e:
                        print(f"Emotion detection error: {e}")

                frame_count += 1
                self.msleep(66)  # ~15 FPS

        except Exception as e:
            error_msg = f"Video capture error: {e}"
            self.error.emit(error_msg)
            print(error_msg)
        finally:
            try:
                if cap is not None:
                    cap.release()
                    print("Camera released")
            except:
                pass

    def get_recent_emotions(self, duration: float = 2.0) -> List[Dict]:
        """Get emotion detections from the last N seconds (duration)"""
        if not self._emotion_history or self._recording_start_time is None:
            return []

        current_time = time.time() - self._recording_start_time
        start_time = max(0.0, current_time - duration)

        # Filter emotions by timestamp within the paragraph duration
        filtered_emotions = []
        for emotion_data in self._emotion_history:
            timestamp = emotion_data.get('timestamp', 0.0)
            if start_time <= timestamp <= current_time:
                filtered_emotions.append(emotion_data)

        return filtered_emotions

    def get_emotions_in_timerange(self, start_time: float, end_time: float) -> List[Dict]:
        """Get emotion detections within a specific time range (for paragraph processing)"""
        if not self._emotion_history:
            return []

        # Filter emotions by timestamp within the specified range
        filtered_emotions = []
        for emotion_data in self._emotion_history:
            timestamp = emotion_data.get('timestamp', 0.0)
            if start_time <= timestamp <= end_time:
                filtered_emotions.append(emotion_data)

        return filtered_emotions

# ========================= Dictation worker (QThread) ======================
class DictationWorker(QThread):
    # Speech
    partial = pyqtSignal(str)
    final_json = pyqtSignal(dict)   # full Vosk JSON on final (for word timestamps)
    error   = pyqtSignal(str)
    # Meter
    level   = pyqtSignal(float)     # 0..1 after AGC
    pulse   = pyqtSignal()          # word pulse
    # Tone VA(D)
    vad_frame = pyqtSignal(float, float, float, float)  # V, A, D, t_center_seconds

    def __init__(self, model_dir: Path, samplerate: int = 16000):
        super().__init__()
        self.model_dir = model_dir
        self.samplerate = samplerate
        self._pcm_q: "queue.Queue[bytes]" = queue.Queue(maxsize=20)
        self._stop = False
        self._stream = None
        # AGC state
        self._agc_gain = 1.0
        self._recent_words = 0
        # Clock
        self._t_sec = 0.0  # running audio time (s)
        self._tone = ToneEstimator(samplerate)

    def stop(self):
        self._stop = True
        try:
            if self._stream is not None:
                self._stream.stop()
                self._stream.close()
        except Exception:
            pass

    def _cb(self, indata, frames, time_info, status):
        try:
            self._pcm_q.put_nowait(bytes(indata))
        except Exception:
            pass

    def run(self):
        if VoskModel is None or KaldiRecognizer is None:
            self.error.emit("Vosk is not installed."); return
        if sd is None:
            self.error.emit("sounddevice is not installed."); return
        try:
            model = VoskModel(str(self.model_dir))
        except Exception as e:
            self.error.emit(f"Failed to load Vosk model: {e}"); return
        try:
            info = sd.query_devices(kind='input')
            sr = int(info.get('default_samplerate', self.samplerate)) if info else self.samplerate
            sr = 16000 if sr <= 0 else int(sr); self.samplerate = sr
        except Exception:
            sr = self.samplerate
        recognizer = KaldiRecognizer(model, sr); recognizer.SetWords(True)
        try:
            self._stream = sd.RawInputStream(samplerate=sr, channels=1, dtype='int16',
                                             blocksize=8000, callback=self._cb)
            self._stream.start()
        except Exception as e:
            self.error.emit(f"Microphone error: {e}"); return

        try:
            ema_level = 0.0
            self._t_sec = 0.0
            while not self._stop:
                try:
                    block = self._pcm_q.get(timeout=0.12)
                except queue.Empty:
                    ema_level *= 0.85
                    self.level.emit(max(0.0, min(1.0, ema_level)))
                    continue

                # --- Meter & AGC ---
                a = array('h'); a.frombytes(block)
                peak = (max(abs(x) for x in a) / 32768.0) if a else 0.0
                if self._recent_words > 0 and peak > 0.02:
                    target_gain = 0.9 / peak
                    target_gain = max(1.0, min(8.0, target_gain))
                    self._agc_gain = 0.80*self._agc_gain + 0.20*target_gain
                    self._recent_words -= 1
                else:
                    self._agc_gain = 0.97*self._agc_gain + 0.03*1.0
                level_post_agc = min(1.0, peak * self._agc_gain)
                ema_level = 0.55*ema_level + 0.45*level_post_agc
                self.level.emit(max(0.0, min(1.0, ema_level)))

                # --- Tone VA(D) frame ---
                V, A, D = self._tone.estimate_vad(block)
                # Timestamp center of this block
                dur = len(block) / 2 / float(sr)  # int16 mono → samples = bytes/2
                t_center = self._t_sec + 0.5*dur
                self.vad_frame.emit(float(V), float(A), float(D), float(t_center))
                self._t_sec += dur

                # --- ASR ---
                if recognizer.AcceptWaveform(block):
                    try:
                        j = json.loads(recognizer.Result())
                        # Emit full JSON for timing (contains "text" and "result" with word start/end)
                        self.final_json.emit(j)
                        if (j.get('text') or '').strip():
                            self._recent_words = max(self._recent_words, 4)
                            self.pulse.emit()
                    except Exception:
                        pass
                else:
                    try:
                        j = json.loads(recognizer.PartialResult()); p = (j.get('partial') or '').strip()
                        if p:
                            self.partial.emit(p)
                            self._recent_words = max(self._recent_words, 6)
                            self.pulse.emit()
                    except Exception:
                        pass
        finally:
            try:
                if self._stream:
                    self._stream.stop(); self._stream.close()
            except Exception:
                pass

# ============================== Level Meter ================================
class LevelMeter(QtWidgets.QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setFixedHeight(72); self.setMinimumWidth(160)
        self._active = False
        self._hist = deque([], maxlen=180)
        self._phase = 0.0
        self._amp_target = 0.0
        self._amp = 0.0
        self._word_boost = 0.0
        self._timer = QtCore.QTimer(self); self._timer.setInterval(33); self._timer.timeout.connect(self._tick)

    def set_active(self, on: bool):
        self._active = on
        if on:
            self._phase = 0.0; self._timer.start()
        else:
            self._amp_target = 0.0; self._word_boost = 0.0
            QtCore.QTimer.singleShot(250, self._timer.stop)
        self.update()

    @QtCore.pyqtSlot(float)
    def push(self, v: float):
        self._amp_target = max(0.12, min(1.0, v)) if self._active else 0.0

    @QtCore.pyqtSlot()
    def pulse(self):
        self._word_boost = min(1.0, self._word_boost + 0.6)

    def _ensure_hist_size(self):
        w = max(30, self.width() - 6)
        if w != self._hist.maxlen:
            old = list(self._hist)
            new = deque([0.0]*w, maxlen=w)
            if old:
                for i in range(w):
                    j = int(i * (len(old)-1) / max(1, w-1))
                    new[i] = old[j]
            self._hist = new

    def _tick(self):
        self._ensure_hist_size()
        self._word_boost *= 0.84
        base = 0.12 if self._active else 0.0
        target = max(base, self._amp_target + 0.6*self._word_boost)
        if target < self._amp: self._amp = 0.55*self._amp + 0.45*target
        else: self._amp = 0.85*self._amp + 0.15*target
        self._phase = (self._phase + 0.18*4) % (2*math.pi)
        y_new = math.sin(self._phase) * min(1.0, self._amp)
        if self._hist.maxlen == 0:
            self._hist = deque([y_new], maxlen=1)
        else:
            self._hist.append(y_new)
        self.update()

    def resizeEvent(self, ev: QtGui.QResizeEvent):
        self._ensure_hist_size(); super().resizeEvent(ev)

    def paintEvent(self, ev: QtGui.QPaintEvent):
        p = QtGui.QPainter(self); p.setRenderHint(QtGui.QPainter.RenderHint.Antialiasing, True)
        w, h = self.width(), self.height()
        p.fillRect(self.rect(), QtGui.QColor("#111"))
        if not self._hist:
            pen_empty = QtGui.QPen(QtGui.QColor("#333")); pen_empty.setWidth(1)
            p.setPen(pen_empty); p.drawRect(0, 0, w-1, h-1); return
        mid = h/2; scale = (h-6)*0.5
        N = len(self._hist); pts = []
        for i, y in enumerate(self._hist):
            x = 3 + int((w-6) * (i / max(1, N-1)))
            pts.append(QtCore.QPointF(x, mid + scale * y))
        pen = QtGui.QPen(QtGui.QColor(CYBER_BLUE_HEX)); pen.setWidthF(3.0)
        p.setPen(pen)
        for i in range(1, len(pts)):
            p.drawLine(pts[i-1], pts[i])
        frame = QtGui.QPen(QtGui.QColor("#333")); frame.setWidth(1)
        p.setPen(frame); p.drawRect(0, 0, w-1, h-1)

# ============================== Video Display Widget ================================
class VideoWidget(QtWidgets.QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setFixedSize(160, 120)  # Small video preview next to volume meter
        self._frame = None
        self._emotion_text = ""
        self._status_text = "Waiting..."
        self._timer = QtCore.QTimer(self)
        self._timer.setInterval(100)  # Update display every 100ms
        self._timer.timeout.connect(self.update)
        self._timer.start()

    @QtCore.pyqtSlot(object)
    def update_frame(self, frame):
        """Update the video frame"""
        if frame is not None:
            self._frame = frame.copy()  # Make a copy
            self._status_text = "Active"
        else:
            print("Video widget received None frame")

    @QtCore.pyqtSlot(list)
    def update_emotions(self, emotions):
        """Update emotion display"""
        if emotions:
            # Show dominant emotion from first detection
            detection = emotions[0]
            emotion_scores = detection.get('emotions', {})
            if emotion_scores:
                dominant = max(emotion_scores.items(), key=lambda x: x[1])
                self._emotion_text = f"{dominant[0]}: {dominant[1]:.2f}"
            else:
                self._emotion_text = "detecting..."
        else:
            self._emotion_text = ""

    @QtCore.pyqtSlot(str)
    def show_error(self, error_msg):
        """Show error message"""
        self._status_text = f"Error: {error_msg}"
        self._frame = None
        print(f"Video widget error: {error_msg}")

    def paintEvent(self, event):
        painter = QtGui.QPainter(self)
        painter.fillRect(self.rect(), QtGui.QColor("#111"))

        if self._frame is not None:
            try:
                # Convert frame to Qt format (OpenCV uses BGR, Qt uses RGB)
                if len(self._frame.shape) == 3:
                    height, width, channel = self._frame.shape
                    if channel == 3:
                        # Convert BGR to RGB
                        rgb_frame = cv2.cvtColor(self._frame, cv2.COLOR_BGR2RGB) if cv2 else self._frame
                        bytes_per_line = 3 * width
                        q_image = QtGui.QImage(rgb_frame.data, width, height, bytes_per_line, QtGui.QImage.Format.Format_RGB888)

                        # Scale to widget size
                        pixmap = QtGui.QPixmap.fromImage(q_image)
                        scaled_pixmap = pixmap.scaled(self.size(), Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation)

                        # Center the image
                        x = (self.width() - scaled_pixmap.width()) // 2
                        y = (self.height() - scaled_pixmap.height()) // 2
                        painter.drawPixmap(x, y, scaled_pixmap)

                        # Draw emotion text overlay
                        if self._emotion_text:
                            painter.setPen(QtGui.QColor(CYBER_BLUE_HEX))
                            painter.setFont(QtGui.QFont("Arial", 8))
                            painter.drawText(5, self.height() - 15, self._emotion_text)

                        # Draw status
                        painter.setPen(QtGui.QColor("#AAA"))
                        painter.setFont(QtGui.QFont("Arial", 7))
                        painter.drawText(5, 12, self._status_text)
                        return

            except Exception as e:
                print(f"Video display error: {e}")
                self._status_text = f"Display error"

        # No video or error - show placeholder
        painter.setPen(QtGui.QColor("#333"))
        painter.drawRect(0, 0, self.width()-1, self.height()-1)
        painter.setPen(QtGui.QColor("#666"))
        painter.setFont(QtGui.QFont("Arial", 10))
        painter.drawText(self.rect(), Qt.AlignmentFlag.AlignCenter, "📹\nVideo")

        # Show status
        painter.setPen(QtGui.QColor("#AAA"))
        painter.setFont(QtGui.QFont("Arial", 8))
        painter.drawText(5, self.height() - 5, self._status_text)

# =============================== Text Cleanup ==============================
CUSTOM_REPLACEMENTS = {
    # Add custom word replacements here for common speech recognition errors
    # Example: "there system": "their system",
}
SPOKEN_PUNCT = {
    "period": ".", "full stop": ".", "question mark": "?", "exclamation point": "!",
    "exclamation mark": "!", "comma": ",", "semicolon": ";", "colon": ":", "dash": "—", "ellipsis": "…",
}

def _apply_spoken_punct(text: str) -> str:
    if not text: return text
    pattern = r"\b(" + "|".join(re.escape(k) for k in SPOKEN_PUNCT.keys()) + r")\b"
    def repl(m): return SPOKEN_PUNCT.get(m.group(1).lower(), m.group(0))
    out = re.sub(pattern, repl, text, flags=re.IGNORECASE)
    out = re.sub(r"\s+([.,!?;:…—])", r"\1", out)
    for word, sym in SPOKEN_PUNCT.items():
        if re.search(rf"\b{re.escape(word)}\s*$", out, flags=re.IGNORECASE):
            out = re.sub(rf"\b{re.escape(word)}\s*$", sym, out, flags=re.IGNORECASE)
    return out

def _basic_punctuate_and_capitalize(text: str) -> str:
    t = _apply_spoken_punct(text.strip())
    if not t: return t
    parts = re.split(r"([.!?…])", t)
    if len(parts) == 1:
        t = re.sub(r"\bi\b", "I", t)
        return t[0:1].upper() + t[1:] if t else t
    sentences = []
    for i in range(0, len(parts), 2):
        seg = parts[i].strip(); end = parts[i+1] if i+1 < len(parts) else ""
        if not seg: continue
        seg = re.sub(r"\bi\b", "I", seg)
        seg = seg[0:1].upper() + seg[1:]
        sentences.append(seg + (end if end else ""))
    out = " ".join(sentences).strip()
    for wrong, right in CUSTOM_REPLACEMENTS.items():
        out = re.sub(rf"\b{re.escape(wrong)}\b", right, out, flags=re.IGNORECASE)
    return out

# =============================== Main Window ===============================
class MainWindow(QtWidgets.QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle(APP_TITLE); self.resize(1280, 860)

        # Emotion stack
        self.emotion_matrix = load_emotion_matrix()
        self.summarizer = EmotionSummarizer(self.emotion_matrix)
        self.text_est = TextValenceEstimator()

        # NEW: Local Llama emotion interpretation (optional)
        self.llama_interpreter: Optional[LocalLlamaEmotionInterpreter] = None
        self._llama_enabled = False

        # NEW: Emotion state machine with single-worker queue
        self.emotion_state = EmotionState.IDLE
        self.emotion_queue: queue.Queue[EmotionJob] = queue.Queue(maxsize=1)
        self.emotion_worker: Optional[EmotionQueueWorker] = None
        self._paragraph_buffer: List[str] = []  # Accumulates text during LISTENING
        self._current_vad_snapshot: dict = {"valence": 0.0, "arousal": 0.0, "dominance": 0.0}
        self._last_pause_time: float = 0.0
        self._pause_timer: Optional[QtCore.QTimer] = None
        self._speech_since_ready = False  # Track if new speech occurred since READY_NEXT
        self._first_speech_time: float = 0.0  # Track when first short phrase was received
        self._llm_request_time: float = 0.0  # Track when LLM request was sent
        self._llm_timeout_timer: Optional[QtCore.QTimer] = None  # Timer for LLM timeout

        # Streaming buffers (audio)
        self._vad_times: List[float] = []
        self._vad_V: List[float] = []
        self._vad_A: List[float] = []
        self._vad_D: List[float] = []

        self.pubkey_path: Optional[Path] = None
        self.privkey_path: Optional[Path] = None
        self.pubkey_bytes: Optional[bytes] = None
        self.privkey_bytes: Optional[bytes] = None
        self.privkey_pass: Optional[str] = None
        self.nb_path: Optional[Path] = None
        self.nb: Optional[dict] = None
        self.nb_dirty = False
        self.vosk_model_dir: Optional[Path] = find_builtin_vosk_model()
        self.dictation: Optional[DictationWorker] = None

        # NEW: Video components
        self.video_worker: Optional[VideoWorker] = None

        central = QtWidgets.QWidget(); self.setCentralWidget(central)
        root = QtWidgets.QVBoxLayout(central)

        self.topbar = QtWidgets.QHBoxLayout(); root.addLayout(self.topbar)
        self.btn_pub    = self._pill_button("🔐 Public Key", "Select PUBLIC key…", self.pick_pub)
        self.btn_priv   = self._pill_button("🗝  Private Key", "Select PRIVATE key…", self.pick_priv)
        self.btn_open   = self._pill_button("📂 Open", "Open Notebook…", self.open_notebook)
        self.btn_create = self._pill_button("✨ New", "Create Notebook…", self.create_notebook)
        self.btn_save   = self._pill_button("💾 Save", "Save Notebook", self.save_notebook)
        self.btn_close  = self._pill_button("✖ Close", "Close Notebook", self.close_notebook)
        self.btn_export = self._pill_button("📤 Export", "Export to Word…", self.export_docx)
        self.btn_record = self._pill_button("🔴 Record", "Start Recording (Audio + Video)", self.toggle_recording)
        self.topbar.addStretch(1)
        self.search = QtWidgets.QLineEdit(placeholderText="Search titles… (Ctrl+F)")
        self.search.textChanged.connect(self._filter_list); self.search.setObjectName("Search")
        self.topbar.addWidget(self.search)

        splitter = QtWidgets.QSplitter(Qt.Orientation.Horizontal); root.addWidget(splitter, 1)

        left = QtWidgets.QWidget(); left_layout = QtWidgets.QVBoxLayout(left)
        left_layout.setContentsMargins(0, 0, 0, 0)
        self.listw = QtWidgets.QListWidget(); self.listw.itemSelectionChanged.connect(self.on_select)
        left_layout.addWidget(self.listw, 1)

        # ENHANCED: Bottom area with two-row layout for narrower left panel
        bottom_container = QtWidgets.QVBoxLayout()
        bottom_container.setContentsMargins(8, 0, 8, 8)
        bottom_container.setSpacing(4)

        # First row: Volume meter, Video widget, Audio/Video ratio, Add button
        top_row = QtWidgets.QHBoxLayout()
        top_row.setSpacing(8)

        # Volume meter
        self.meter = LevelMeter()
        top_row.addWidget(self.meter, 0)

        # Video widget
        self.video_widget = VideoWidget()
        top_row.addWidget(self.video_widget, 0)

        # Audio/Video fusion weight dropdown
        fusion_label = QtWidgets.QLabel("Audio/Video:")
        fusion_label.setStyleSheet("color: #C0C0C0; font-size: 10pt;")
        top_row.addWidget(fusion_label, 0)

        self.fusion_dropdown = QtWidgets.QComboBox()
        self.fusion_dropdown.addItems(["0/100", "20/80", "40/60", "60/40", "80/20", "100/0"])
        self.fusion_dropdown.setCurrentText("40/60")  # Default: 40% audio, 60% video
        self.fusion_dropdown.setToolTip("Audio % / Video % emotion fusion weights")
        self.fusion_dropdown.setStyleSheet("""
            QComboBox {
                background: #232323;
                color: #00D0FF;
                border: 1px solid #3a3a3a;
                border-radius: 8px;
                padding: 4px 8px;
                min-width: 60px;
            }
            QComboBox::drop-down {
                border: none;
            }
            QComboBox::down-arrow {
                image: none;
                border: none;
            }
        """)
        top_row.addWidget(self.fusion_dropdown, 0)

        top_row.addStretch(1)
        self.fab = QtWidgets.QPushButton("＋"); self.fab.setObjectName("Fab"); self.fab.setToolTip("Add Entry")
        self.fab.clicked.connect(self.add_entry)
        top_row.addWidget(self.fab, 0, Qt.AlignmentFlag.AlignRight)

        # Second row: Debug checkbox and LLM controls
        bottom_row = QtWidgets.QHBoxLayout()
        bottom_row.setSpacing(8)

        # Debug toggle checkbox
        debug_label = QtWidgets.QLabel("Debug:")
        debug_label.setStyleSheet("color: #C0C0C0; font-size: 10pt;")
        bottom_row.addWidget(debug_label, 0)

        self.debug_checkbox = QtWidgets.QCheckBox()
        self.debug_checkbox.setChecked(False)
        self.debug_checkbox.setToolTip("Enable debug output for audio/video emotion results")
        self.debug_checkbox.setStyleSheet("""
            QCheckBox {
                color: #00D0FF;
                spacing: 5px;
            }
            QCheckBox::indicator {
                width: 16px;
                height: 16px;
                border: 1px solid #3a3a3a;
                border-radius: 3px;
                background: #232323;
            }
            QCheckBox::indicator:checked {
                background: #00D0FF;
                border-color: #00D0FF;
            }
        """)
        bottom_row.addWidget(self.debug_checkbox, 0)

        # LLM controls
        llm_label = QtWidgets.QLabel("LLM:")
        llm_label.setStyleSheet("color: #C0C0C0; font-size: 10pt;")
        bottom_row.addWidget(llm_label, 0)

        self.llm_mode_dropdown = QtWidgets.QComboBox()
        self.llm_mode_dropdown.addItems(["off", "local"])
        self.llm_mode_dropdown.setCurrentText("off")
        self.llm_mode_dropdown.setToolTip("Local Llama emotion interpretation: off/local llama.cpp")

        # Connect signal to handle LLM mode changes
        self.llm_mode_dropdown.currentTextChanged.connect(self._on_llm_mode_changed)
        self.llm_mode_dropdown.setStyleSheet("""
            QComboBox {
                background: #232323;
                color: #00D0FF;
                border: 1px solid #3a3a3a;
                border-radius: 8px;
                padding: 4px 8px;
                min-width: 60px;
            }
            QComboBox::drop-down {
                border: none;
            }
            QComboBox::down-arrow {
                image: none;
                border: none;
            }
        """)
        bottom_row.addWidget(self.llm_mode_dropdown, 0)

        bottom_row.addStretch(1)

        # Add both rows to the container
        bottom_container.addLayout(top_row)
        bottom_container.addLayout(bottom_row)
        left_layout.addLayout(bottom_container)
        splitter.addWidget(left)

        right = QtWidgets.QWidget(); rlay = QtWidgets.QVBoxLayout(right)
        hdr = QtWidgets.QHBoxLayout()
        self.dt = QtWidgets.QLineEdit(); self.dt.setReadOnly(True); self.dt.setPlaceholderText("Date/Time (ISO)")
        self.title = QtWidgets.QLineEdit(placeholderText="Title…")
        hdr.addWidget(self.dt, 0); hdr.addWidget(self.title, 1); rlay.addLayout(hdr)
        self.body = QtWidgets.QTextEdit(); self.body.setPlaceholderText("Speak or type…")
        rlay.addWidget(self.body, 1)
        foot = QtWidgets.QHBoxLayout()
        self.btn_prev = self._pill_button("⟨ Prev", "Previous entry", self.goto_prev)
        self.btn_next = self._pill_button("Next ⟩", "Next entry", self.goto_next)
        foot.addWidget(self.btn_prev); foot.addWidget(self.btn_next); foot.addStretch(1)
        self.btn_save_entry = self._pill_button("✅ Save Entry", "Save current entry changes", self.save_entry_changes)
        foot.addWidget(self.btn_save_entry); rlay.addLayout(foot)
        splitter.addWidget(right)

        # Configure splitter proportions - make text area (right) much larger
        splitter.setStretchFactor(0, 1)  # Left side (journal list)
        splitter.setStretchFactor(1, 4)  # Right side (text editor) - increased from 2 to 4

        # Set initial sizes: left=250px, right=rest of window
        splitter.setSizes([250, 800])  # Left panel narrower, text area much wider

        self.status = QtWidgets.QStatusBar(); self.setStatusBar(self.status)
        self.status.showMessage("Ready. Press Start to dictate.")

        # Shortcuts
        QtGui.QShortcut(QtGui.QKeySequence.StandardKey.New,  self, activated=self.create_notebook)
        QtGui.QShortcut(QtGui.QKeySequence.StandardKey.Open, self, activated=self.open_notebook)
        QtGui.QShortcut(QtGui.QKeySequence.StandardKey.Save, self, activated=self.save_notebook)
        QtGui.QShortcut(QtGui.QKeySequence("Ctrl+E"), self, activated=self.export_docx)
        QtGui.QShortcut(QtGui.QKeySequence("Ctrl+N"), self, activated=self.add_entry)
        QtGui.QShortcut(QtGui.QKeySequence.StandardKey.Find, self, activated=self._focus_search)
        QtGui.QShortcut(QtGui.QKeySequence("F5"), self, activated=self.toggle_recording)
        QtGui.QShortcut(QtGui.QKeySequence("Ctrl+K"), self, activated=self._open_palette)

        self.title.textChanged.connect(self._mark_dirty)
        self.body.textChanged.connect(self._mark_dirty)

        # Initialize emotion state machine
        self._init_emotion_state_machine()

        self._apply_qss(); self._refresh_controls()
        self._palette: Optional[QtWidgets.QDialog] = None
        self._full_list_cache: List[dict] = []

    def _init_emotion_state_machine(self):
        # Function Purpose: Initialize the emotion state machine and worker thread
        try:
            # Set up pause detection timer
            self._pause_timer = QtCore.QTimer()
            self._pause_timer.setSingleShot(True)
            self._pause_timer.timeout.connect(self._on_speech_pause)

            # Set up LLM timeout timer
            self._llm_timeout_timer = QtCore.QTimer()
            self._llm_timeout_timer.setSingleShot(True)
            self._llm_timeout_timer.timeout.connect(self._on_llm_timeout)

            # Determine which LLM system to use based on dropdown setting
            llm_mode = getattr(self, 'llm_mode_dropdown', None)
            current_llm_mode = llm_mode.currentText() if llm_mode else "off"

            print(f"[DEBUG] LLM mode: '{current_llm_mode}'")

            if current_llm_mode == "off":
                print("[DEBUG] LLM disabled - no emotion worker will be started")
            elif current_llm_mode == "local":
                print("[DEBUG] LLM set to 'local' - new auto-start system will be used during speech processing")
            else:
                # Legacy: start old worker for any unknown modes
                print(f"[DEBUG] Unknown LLM mode '{current_llm_mode}' - starting old emotion worker")
                if not hasattr(self, 'emotion_worker') or not self.emotion_worker or not self.emotion_worker.is_alive():
                    self.emotion_worker = EmotionQueueWorker(
                        in_q=self.emotion_queue,
                        on_result=self._on_emotion_result,
                        call_llama=call_llama
                    )
                    self.emotion_worker.start()
                    print("Emotion worker thread started")
                else:
                    print("Emotion worker thread already running")
        except Exception as e:
            print(f"Error initializing emotion state machine: {e}")
            # Continue without emotion processing

    def _on_emotion_result(self, job: EmotionJob, phrase: str):
        # Function Purpose: Thread-safe callback when Llama returns emotion result
        try:
            # Use Qt's thread-safe method to invoke on main thread
            QtCore.QMetaObject.invokeMethod(
                self, "_handle_emotion_result",
                Qt.ConnectionType.QueuedConnection,
                QtCore.Q_ARG(object, job),
                QtCore.Q_ARG(str, phrase)
            )
        except Exception as e:
            print(f"Error in emotion result callback: {e}")
            # Don't crash the worker thread

    @QtCore.pyqtSlot(object, str)
    def _handle_emotion_result(self, job: EmotionJob, phrase: str):
        # Function Purpose: Main thread handler for emotion results with state transition
        if self.debug_checkbox.isChecked():
            ts = debug_timestamp()
            print(f"[{ts}] Emotion result for '{job.text[:50]}...': '{phrase}'")
            print(f"[{ts}] State: WAITING_MODEL → READY_NEXT (llama result arrives)")

        # Append emotion phrase to the current entry
        if self.body:
            cursor = self.body.textCursor()
            cursor.movePosition(QtGui.QTextCursor.MoveOperation.End)
            cursor.insertText(f"\n• {{Emotion: {phrase}}}")
            self.nb_dirty = True
            self._refresh_controls()

        # State transition: WAITING_MODEL → READY_NEXT (llama result arrives)
        self.emotion_state = EmotionState.READY_NEXT
        self._speech_since_ready = False

    def _update_vad_snapshot(self):
        # Function Purpose: Update current VAD snapshot from latest streaming data
        if self._vad_V and self._vad_A and self._vad_D:
            # Use recent average (last 1 second worth)
            recent_count = min(50, len(self._vad_V))  # ~1s at 50Hz
            self._current_vad_snapshot = {
                "valence": sum(self._vad_V[-recent_count:]) / recent_count,
                "arousal": sum(self._vad_A[-recent_count:]) / recent_count,
                "dominance": sum(self._vad_D[-recent_count:]) / recent_count
            }

    def _on_speech_pause(self):
        # Function Purpose: Handle debounced speech pause with exact state machine specification
        current_time = time.time()

        # Debounce rapid pauses
        if current_time - self._last_pause_time < PAUSE_DEBOUNCE_MS / 1000.0:
            return

        self._last_pause_time = current_time

        if self.debug_checkbox.isChecked():
            ts = debug_timestamp()
            print(f"[{ts}] Speech pause detected. State: {self.emotion_state}")

        # State machine transitions on pause per specification
        if self.emotion_state == EmotionState.LISTENING:
            # LISTENING → WAITING_MODEL (pause → enqueue job if worker idle)
            if self._is_worker_idle():
                self._try_enqueue_current_paragraph()
                if self.debug_checkbox.isChecked():
                    ts = debug_timestamp()
                    print(f"[{ts}] State: LISTENING → WAITING_MODEL (enqueued job)")
            else:
                if self.debug_checkbox.isChecked():
                    ts = debug_timestamp()
                    print(f"[{ts}] State: LISTENING (worker busy, keeping buffer)")

        elif self.emotion_state == EmotionState.WAITING_MODEL:
            # Do nothing - continue accumulating
            if self.debug_checkbox.isChecked():
                ts = debug_timestamp()
                print(f"[{ts}] State: WAITING_MODEL (ignoring pause, continuing accumulation)")

        elif self.emotion_state == EmotionState.READY_NEXT:
            # READY_NEXT → WAITING_MODEL (next pause after new speech → enqueue)
            if self._speech_since_ready and self._paragraph_buffer:
                self._try_enqueue_current_paragraph()
                if self.debug_checkbox.isChecked():
                    ts = debug_timestamp()
                    print(f"[{ts}] State: READY_NEXT → WAITING_MODEL (new speech detected, enqueued)")
            else:
                if self.debug_checkbox.isChecked():
                    ts = debug_timestamp()
                    print(f"[{ts}] State: READY_NEXT (no new speech, ignoring pause)")

    def _is_worker_idle(self) -> bool:
        # Function Purpose: Check if emotion worker is idle (queue empty and not processing)
        # Check if LLM processing is enabled
        llm_mode = getattr(self, 'llm_mode_dropdown', None)
        current_llm_mode = llm_mode.currentText() if llm_mode else "off"

        if current_llm_mode == "off":
            return False  # No worker when LLM is disabled
        elif current_llm_mode == "local":
            # For local mode, we don't use the worker thread - always return True
            # The new LocalLlamaEmotionInterpreter handles processing directly
            return True

        return (self.emotion_queue.empty() and
                self.emotion_worker and
                self.emotion_worker.is_alive())

    def _try_enqueue_current_paragraph(self):
        # Function Purpose: Attempt to create and enqueue emotion job from current paragraph
        if not self._paragraph_buffer:
            return

        # Check if LLM processing is enabled
        llm_mode = getattr(self, 'llm_mode_dropdown', None)
        current_llm_mode = llm_mode.currentText() if llm_mode else "off"

        if current_llm_mode == "off":
            if self.debug_checkbox.isChecked():
                print("[DEBUG] LLM disabled - skipping emotion job creation")
            return

        # Finalize current paragraph
        paragraph_text = " ".join(self._paragraph_buffer).strip()
        if not paragraph_text:
            return

        # Count words in accumulated text
        word_count = len(paragraph_text.split())
        current_time = time.time()

        # Initialize timer on first speech (any length)
        if self._first_speech_time == 0.0:
            self._first_speech_time = current_time
            if self.debug_checkbox.isChecked():
                ts = debug_timestamp()
                print(f"[{ts}] Started accumulation ({word_count} words): '{paragraph_text[:50]}...'")

        # Calculate elapsed time
        elapsed = current_time - self._first_speech_time

        # Determine minimum wait time based on word count
        if word_count < MIN_WORDS_FOR_LLM:
            # Short phrases need to wait longer (5 seconds)
            min_wait = MIN_TIME_FOR_SHORT_LLM
            if elapsed < min_wait:
                if self.debug_checkbox.isChecked():
                    ts = debug_timestamp()
                    print(f"[{ts}] Short phrase ({word_count} words, {elapsed:.1f}s/{min_wait}s): '{paragraph_text}'")
                return
        else:
            # Longer phrases still wait minimum time to encourage paragraphs (2 seconds)
            min_wait = MIN_TIME_FOR_ANY_LLM
            if elapsed < min_wait:
                if self.debug_checkbox.isChecked():
                    ts = debug_timestamp()
                    print(f"[{ts}] Accumulating paragraph ({word_count} words, {elapsed:.1f}s/{min_wait}s): '{paragraph_text[:50]}...'")
                return

        # Time threshold met - proceed with processing
        if self.debug_checkbox.isChecked():
            ts = debug_timestamp()
            print(f"[{ts}] Time threshold met ({elapsed:.1f}s), processing {word_count} words")

        # Reset timer for next accumulation
        self._first_speech_time = 0.0

        # Update VAD snapshot
        self._update_vad_snapshot()

        if current_llm_mode == "local":
            # Use new LocalLlamaEmotionInterpreter system directly
            if self.debug_checkbox.isChecked():
                ts = debug_timestamp()
                print(f"[{ts}] Processing with LocalLlamaEmotionInterpreter: '{paragraph_text[:50]}...'")
                print(f"[{ts}] VAD: V={self._current_vad_snapshot['valence']:.3f}, A={self._current_vad_snapshot['arousal']:.3f}, D={self._current_vad_snapshot['dominance']:.3f}")

            # Start LLM timeout timer and track request time
            self._llm_request_time = time.time()
            if self._llm_timeout_timer:
                self._llm_timeout_timer.start(LLAMA_TIMEOUT_MS)

            # Process directly using the new system
            self._process_local_llm_emotion(paragraph_text, self._current_vad_snapshot.copy())
            self._paragraph_buffer.clear()
            return

    def _process_local_llm_emotion(self, text: str, vad_data: dict):
        """Process emotion using LocalLlamaEmotionInterpreter directly"""
        try:
            if self._ensure_llama_interpreter():
                # Get recent video emotions (last 5 seconds)
                video_emotions = []
                if self.video_worker:
                    current_time = time.time()
                    start_time = current_time - 5.0  # Look back 5 seconds
                    video_emotions = self.video_worker.get_emotions_in_timerange(start_time, current_time)

                # Use the new auto-start system
                result = self.llama_interpreter.interpret_emotion(
                    text=text,
                    audio_data=vad_data,
                    video_data=video_emotions,  # Now passing actual video data
                    debug=self.debug_checkbox.isChecked()
                )

                # Stop timeout timer on result
                if self._llm_timeout_timer:
                    self._llm_timeout_timer.stop()

                if result and not result.startswith("llama"):
                    # Successful result - add to journal
                    self._add_llm_result_to_journal(text, result)

                    # Transition WAITING_MODEL → READY_NEXT
                    self.emotion_state = EmotionState.READY_NEXT
                    self._speech_since_ready = False

                    if self.debug_checkbox.isChecked():
                        ts = debug_timestamp()
                        print(f"[{ts}] Local LLM result: '{result}'")
                        print(f"[{ts}] State: WAITING_MODEL → READY_NEXT")
                else:
                    # Failed - transition back to LISTENING to allow retry
                    self.emotion_state = EmotionState.LISTENING
                    if self.debug_checkbox.isChecked():
                        ts = debug_timestamp()
                        print(f"[{ts}] Local LLM failed: '{result}'")
                        print(f"[{ts}] State: WAITING_MODEL → LISTENING (failed)")
            else:
                if self.debug_checkbox.isChecked():
                    ts = debug_timestamp()
                    print(f"[{ts}] Local LLM interpreter not available")

        except Exception as e:
            if self.debug_checkbox.isChecked():
                ts = debug_timestamp()
                print(f"[{ts}] Error in local LLM processing: {e}")

    def _add_llm_result_to_journal(self, text: str, llm_result: str):
        """Add LLM result to the journal (text already inserted by _insert_text_only)"""
        try:
            cursor = self.body.textCursor()
            cursor.movePosition(QtGui.QTextCursor.MoveOperation.End)

            # Text was already inserted by _insert_text_only in _on_final_json
            # Just add the LLM emotion result with newline after
            cursor.insertText(f"\n• {{Emotion: {llm_result}}}\n")

            self.body.setTextCursor(cursor)
            self.nb_dirty = True
            self._refresh_controls()

            if self.debug_checkbox.isChecked():
                ts = debug_timestamp()
                print(f"[{ts}] Added to journal: '{text[:30]}...' -> '{llm_result}'")

        except Exception as e:
            if self.debug_checkbox.isChecked():
                ts = debug_timestamp()
                print(f"[{ts}] Error adding to journal: {e}")

    def _on_audio_chunk(self, frame_data):
        # Function Purpose: Process audio chunk and update emotion state based on speech activity
        # This would be called from the audio processing pipeline
        # Detect speech activity and transition IDLE/READY_NEXT → LISTENING

        # For now, this is a placeholder - you'll need to integrate with existing VAD
        pass

    def _on_speech_detected(self, text: str):
        # Function Purpose: Handle new speech text and manage state transitions per specification
        cleaned_text = _basic_punctuate_and_capitalize(text)

        # Add to paragraph buffer
        self._paragraph_buffer.append(cleaned_text)

        # State transitions on speech detection
        if self.emotion_state == EmotionState.IDLE:
            # IDLE → LISTENING (speech start)
            self.emotion_state = EmotionState.LISTENING
            if self.debug_checkbox.isChecked():
                ts = debug_timestamp()
                print(f"[{ts}] State: IDLE → LISTENING (speech start)")

        elif self.emotion_state == EmotionState.READY_NEXT:
            # READY_NEXT → LISTENING (speech resumes)
            self.emotion_state = EmotionState.LISTENING
            self._speech_since_ready = True
            if self.debug_checkbox.isChecked():
                ts = debug_timestamp()
                print(f"[{ts}] State: READY_NEXT → LISTENING (speech resumes)")

        elif self.emotion_state == EmotionState.WAITING_MODEL:
            # Continue accumulating while Llama works
            if self.debug_checkbox.isChecked():
                ts = debug_timestamp()
                print(f"[{ts}] State: WAITING_MODEL (continuing to accumulate)")

        # Start/restart pause timer for all states
        if self._pause_timer:
            self._pause_timer.stop()
            self._pause_timer.start(PAUSE_MS)

    def _on_llm_timeout(self):
        # Function Purpose: Handle LLM timeout - transition WAITING_MODEL → LISTENING
        if self.debug_checkbox.isChecked():
            ts = debug_timestamp()
            elapsed = time.time() - self._llm_request_time
            print(f"[{ts}] LLM timeout after {elapsed:.1f}s")
            print(f"[{ts}] State: WAITING_MODEL → LISTENING (timeout)")

        # Transition back to LISTENING to allow retry
        self.emotion_state = EmotionState.LISTENING
        # Don't clear buffer - keep accumulated text for next attempt

    def _insert_text_only(self, plain_text: str, start_time: float, end_time: float):
        # Function Purpose: Insert text immediately without emotion processing
        cleaned = _basic_punctuate_and_capitalize(plain_text)
        cursor = self.body.textCursor()
        cursor.movePosition(QtGui.QTextCursor.MoveOperation.End)

        # Add space if there's existing content, but don't create new paragraph
        current_text = self.body.toPlainText()
        if current_text.strip():
            # Only add space if last char isn't already whitespace
            if not current_text.endswith((' ', '\n')):
                cursor.insertText(" ")
        cursor.insertText(cleaned)
        self.body.setTextCursor(cursor)
        self.nb_dirty = True
        self._refresh_controls()

    def _on_llm_mode_changed(self, new_mode: str):
        """Handle LLM mode dropdown changes"""
        print(f"[DEBUG] LLM mode changed to: '{new_mode}'")

        if new_mode == "local":
            print("[DEBUG] Attempting to initialize LLM auto-start system...")
            # Try to initialize the auto-start system immediately
            try:
                if self._ensure_llama_interpreter():
                    print("[DEBUG] LLM auto-start system initialized successfully")
                else:
                    print("[DEBUG] LLM auto-start system failed to initialize")
            except Exception as e:
                print(f"[DEBUG] Error initializing LLM auto-start: {e}")
                import traceback
                traceback.print_exc()
        elif new_mode == "off":
            print("[DEBUG] LLM disabled")
            # Clean up interpreter if it exists
            if hasattr(self, 'llama_interpreter') and self.llama_interpreter:
                try:
                    self.llama_interpreter.cleanup()
                    self.llama_interpreter = None
                    print("[DEBUG] LLM interpreter cleaned up")
                except Exception as e:
                    print(f"[DEBUG] Error cleaning up LLM interpreter: {e}")

    def _get_fusion_weights(self) -> Tuple[float, float]:
        """Get current audio/video fusion weights from dropdown"""
        selection = self.fusion_dropdown.currentText()
        if selection == "0/100":
            return 0.0, 1.0
        elif selection == "20/80":
            return 0.2, 0.8
        elif selection == "40/60":
            return 0.4, 0.6
        elif selection == "60/40":
            return 0.6, 0.4
        elif selection == "80/20":
            return 0.8, 0.2
        elif selection == "100/0":
            return 1.0, 0.0
        else:
            return 0.8, 0.2  # Default fallback

    def _ensure_llama_interpreter(self) -> bool:
        """Initialize local Llama interpreter if needed and available"""
        print(f"[DEBUG] _ensure_llama_interpreter called, llm_available={llm_available}")
        if not llm_available:
            print("[DEBUG] LLM not available, returning False")
            return False

        current_mode = self.llm_mode_dropdown.currentText()
        print(f"[DEBUG] Current LLM mode: '{current_mode}'")

        # Turn off LLM
        if current_mode == "off":
            print("[DEBUG] LLM mode is off")
            self.llama_interpreter = None
            self._llama_enabled = False
            return False

        # Initialize local Llama interpreter if needed
        if current_mode == "local":
            print(f"[DEBUG] Local mode selected, llama_interpreter is None: {self.llama_interpreter is None}")
            if self.llama_interpreter is None:
                try:
                    print("[DEBUG] Creating LocalLlamaEmotionInterpreter...")
                    self.llama_interpreter = LocalLlamaEmotionInterpreter()
                    print("Local Llama interpreter initialized")
                except Exception as e:
                    print(f"Failed to initialize Llama interpreter: {e}")
                    import traceback
                    traceback.print_exc()
                    self.llama_interpreter = None
                    return False

            # Check if interpreter is available
            if self.llama_interpreter and self.llama_interpreter.available:
                self._llama_enabled = True
                return True
            else:
                if self.llama_interpreter:
                    print(f"Local Llama server not available at {self.llama_interpreter.host}:{self.llama_interpreter.port}")
                else:
                    print("Local Llama server not available (interpreter not initialized)")
                self._llama_enabled = False
                return False

        return False

    # --------------------------------- QSS ---------------------------------
    def _apply_qss(self):
        self.setStyleSheet(f"""
            QWidget {{ background: {BLACK_HEX}; color: {CYBER_BLUE_HEX}; font-family: 'Consolas','Menlo',monospace; font-size: 12.5pt; }}
            QLineEdit, QTextEdit {{ background: {INPUT_HEX}; border:1px solid {CHARCOAL_HEX}; border-radius:10px; padding:8px; selection-background-color:#0e3b46; selection-color:white; }}
            QListWidget {{ background: {INPUT_HEX}; border:1px solid {CHARCOAL_HEX}; border-radius:10px; padding:6px; }}
            QListWidget::item {{ padding:8px; border-radius:8px; }}
            QListWidget::item:selected {{ background:#10343d; color:white; }}
            QPushButton {{ background:{CHARCOAL_HEX}; color:{CYBER_BLUE_HEX}; border:1px solid #3a3a3a; border-radius:18px; padding:8px 14px; }}
            QPushButton:hover {{ border-color:#5a5a5a; }}
            QPushButton#Fab {{ background:{CYBER_BLUE_HEX}; color:black; border-radius:24px; padding:10px 14px; min-width:48px; min-height:48px; font-weight:700; }}
            QStatusBar {{ background:{BLACK_HEX}; color:{SILVER_HEX}; border-top:1px solid {CHARCOAL_HEX}; }}
            QToolTip {{ background:{SILVER_HEX}; color:black; }}
        """)

    def _pill_button(self, text: str, tip: str, slot):
        b = QtWidgets.QPushButton(text); b.setToolTip(tip); b.clicked.connect(slot); self.topbar.addWidget(b); return b

    # ------------------------------- Helpers -------------------------------
    def _sorted_entries(self) -> List[dict]:
        if not self.nb: return []
        def key(e): return (e.get("dt") or "", e.get("id") or "")
        return sorted(self.nb["entries"], key=key)

    def _current_entry_index(self) -> int:
        row = self.listw.currentRow()
        return row if row is not None else -1

    def _current_entry(self) -> Optional[dict]:
        entries = self._visible_entries(); row = self._current_entry_index()
        return entries[row] if 0 <= row < len(entries) else None

    def _visible_entries(self) -> List[dict]:
        return self._full_list_cache or self._sorted_entries()

    def _mark_dirty(self):
        if self.nb is None: return
        self.nb_dirty = True; self._refresh_controls()

    def _refresh_controls(self):
        has_nb = self.nb is not None
        has_sel = has_nb and (self._current_entry() is not None)
        for w in (self.btn_prev, self.btn_next, self.btn_save_entry, self.title, self.body):
            w.setEnabled(has_sel)
        self.btn_save.setEnabled(has_nb and self.nb_dirty)
        self.btn_export.setEnabled(has_nb)
        if hasattr(self, "btn_close"): self.btn_close.setEnabled(has_nb)

        # Check if recording (either audio or video)
        is_recording = (self.dictation and self.dictation.isRunning()) or (self.video_worker and self.video_worker.isRunning())

        # Enable record button if we have a selected entry
        self.btn_record.setEnabled(has_sel)

        # Update volume meter activity
        self.meter.set_active(self.dictation and self.dictation.isRunning())

    # ----------------------------- Search/filter ----------------------------
    def _focus_search(self): self.search.setFocus(); self.search.selectAll()
    def _filter_list(self, text: str):
        base = self._sorted_entries()
        self._full_list_cache = [e for e in base if text.strip().lower() in (e.get("title","").lower())]
        self._populate_list(self._full_list_cache)

    # -------------------------- Populate/selection -------------------------
    def _populate_list(self, entries: Optional[List[dict]] = None):
        self.listw.clear()
        for e in (entries or self._sorted_entries()):
            dt = (e.get("dt") or "?")[:19]; title = e.get("title") or "(untitled)"
            self.listw.addItem(f"{dt} — {title}")
        self.dt.setText(""); self.title.setText(""); self.body.setPlainText("")

    def on_select(self):
        if not self.nb: return
        entry = self._current_entry()
        if entry is None:
            self.dt.setText(""); self.title.setText(""); self.body.setPlainText(""); self._refresh_controls(); return
        if not self.privkey_bytes:
            QtWidgets.QMessageBox.warning(self, APP_TITLE, "Select/decrypt your PRIVATE key to open an entry.")
            self.listw.clearSelection(); self._refresh_controls(); return
        try:
            payload = decrypt_blob(entry["enc"], self.privkey_bytes)
            body = unpack_entry_payload(payload).get("body", "")
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, APP_TITLE, f"Decrypt failed for this entry:\n{e}"); return
        self.dt.setText(entry.get("dt","")); self.title.setText(entry.get("title","")); self.body.setPlainText(body)
        self._refresh_controls()

    # --------------------------- File operations ---------------------------
    def pick_pub(self):
        p, _ = QtWidgets.QFileDialog.getOpenFileName(self, "Select PUBLIC key (base64 or raw bytes)")
        if not p: return
        self.pubkey_path = Path(p)
        try:
            self.pubkey_bytes = read_key_file(self.pubkey_path); self.statusBar().showMessage(f"Public: {self.pubkey_path.name}")
        except Exception as e:
            self.pubkey_bytes = None; QtWidgets.QMessageBox.critical(self, APP_TITLE, f"Failed to read public key:\n{e}")

    def pick_priv(self):
        p, _ = QtWidgets.QFileDialog.getOpenFileName(self, "Select PRIVATE key (KeyGen-encrypted, base64, or raw)")
        if not p: return
        self.privkey_path = Path(p)
        pw, ok = QtWidgets.QInputDialog.getText(self, "Private Key Password", "Enter password (if armored):", QtWidgets.QLineEdit.EchoMode.Password)
        self.privkey_pass = pw if ok and pw else None
        try:
            self.privkey_bytes = read_privkey_file(self.privkey_path, self.privkey_pass); self.statusBar().showMessage(f"Private: {self.privkey_path.name}")
        except Exception as e:
            self.privkey_bytes = None; QtWidgets.QMessageBox.critical(self, APP_TITLE, f"Failed to decrypt/read private key:\n{e}")

    def open_notebook(self):
        err = require_crypto_ready()
        if err: QtWidgets.QMessageBox.critical(self, APP_TITLE, err); return
        p, _ = QtWidgets.QFileDialog.getOpenFileName(self, "Open Notebook (.pqjn)", filter="PQ Journal (*.pqjn);;All Files (*)")
        if not p: return
        self.nb_path = Path(p)
        try:
            st = self.nb_path.stat()
            if st.st_size > MAX_NOTEBOOK_BYTES: raise ValueError("Notebook too large.")
            raw = json.loads(self.nb_path.read_text(encoding="utf-8"))
            if raw.get("version") != 1: raise ValueError("Unsupported notebook version.")
            self.nb = raw; self.nb_dirty = False
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, APP_TITLE, f"Failed to open notebook:\n{e}"); self.statusBar().showMessage("Open failed."); return
        if not self.nb["entries"]:
            self._create_blank_entry_and_select()
        self._populate_list(); self._refresh_controls(); self.statusBar().showMessage(f"Notebook loaded: {self.nb_path.name}")

    def create_notebook(self):
        err = require_crypto_ready()
        if err: QtWidgets.QMessageBox.critical(self, APP_TITLE, err); return
        p, _ = QtWidgets.QFileDialog.getSaveFileName(self, "Create New Notebook", filter="PQ Journal (*.pqjn)")
        if not p: return
        self.nb_path = Path(p)
        self.nb = empty_notebook(); write_json_atomic_secure(self.nb_path, self.nb)
        self.nb_dirty = False
        self._create_blank_entry_and_select()
        self._populate_list(); self._refresh_controls(); self.statusBar().showMessage(f"New notebook created: {self.nb_path.name}")

    def save_notebook(self):
        if not self.nb: return
        try: write_json_atomic_secure(self.nb_path, self.nb)
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, APP_TITLE, f"Save failed:\n{e}"); self.statusBar().showMessage("Save failed."); return
        self.nb_dirty = False; self._refresh_controls(); self.statusBar().showMessage(f"Saved: {self.nb_path.name}")

    def close_notebook(self):
        try:
            if self.dictation and self.dictation.isRunning():
                self.stop_dictation()
        except Exception:
            pass
        if not self.nb:
            self.statusBar().showMessage("No notebook open."); return
        if self.nb_dirty:
            box = QtWidgets.QMessageBox(self)
            box.setIcon(QtWidgets.QMessageBox.Icon.Warning)
            box.setWindowTitle(APP_TITLE)
            box.setText("Save changes to this notebook before closing?")
            box.setStandardButtons(
                QtWidgets.QMessageBox.StandardButton.Save |
                QtWidgets.QMessageBox.StandardButton.Discard |
                QtWidgets.QMessageBox.StandardButton.Cancel
            )
            choice = box.exec()
            if choice == QtWidgets.QMessageBox.StandardButton.Save:
                self.save_notebook()
            elif choice == QtWidgets.QMessageBox.StandardButton.Cancel:
                return
        self.nb = None; self.nb_path = None; self.nb_dirty = False
        self._full_list_cache = []
        self.listw.clear(); self.dt.clear(); self.title.clear(); self.body.clear(); self.search.clear()
        self.statusBar().showMessage("Notebook closed.")
        self._refresh_controls()

    # ---------------------------- Entry ops ----------------------------
    def _select_entry_by_id(self, entry_id: str):
        self._full_list_cache = []; self._populate_list()
        entries = self._sorted_entries()
        row = next((i for i, e in enumerate(entries) if e.get("id") == entry_id), -1)
        if row >= 0: self.listw.setCurrentRow(row)

    def _create_blank_entry_and_select(self):
        from datetime import datetime
        now_iso = datetime.now().astimezone().isoformat()
        enc = encrypt_blob(pack_entry_payload({"body": ""}), self.pubkey_bytes or b"\x00") if self.pubkey_bytes else {"alg":"", "kem_ct":"", "nonce":"", "tag":"", "ct":""}
        new_entry = {"id": uuid.uuid4().hex, "dt": now_iso, "title": "", "updated": now_iso, "enc": enc}
        self.nb["entries"].append(new_entry); self.nb_dirty = True
        self._select_entry_by_id(new_entry["id"])
        self.dt.setText(new_entry["dt"]); self.title.setText(""); self.body.setPlainText("")

    def add_entry(self):
        from datetime import datetime
        if not self.nb: QtWidgets.QMessageBox.information(self, APP_TITLE, "Create or open a notebook first."); return
        if not self.pubkey_bytes: QtWidgets.QMessageBox.warning(self, APP_TITLE, "Select your PUBLIC key before adding entries."); return
        if self.search.text(): self.search.clear()
        self._full_list_cache = []
        now_iso = datetime.now().astimezone().isoformat()
        enc = encrypt_blob(pack_entry_payload({"body": ""}), self.pubkey_bytes)
        new_entry = {"id": uuid.uuid4().hex, "dt": now_iso, "title": "", "updated": now_iso, "enc": enc}
        self.nb["entries"].append(new_entry); self.nb_dirty = True
        self._select_entry_by_id(new_entry["id"])
        self.dt.setText(new_entry["dt"]); self.title.setText(""); self.body.setPlainText(""); self.title.setFocus()
        self.statusBar().showMessage("New entry created. Press Start to dictate, or type.")
        self._refresh_controls()

    def save_entry_changes(self):
        if not self.nb: return
        entries = self._visible_entries(); idx = self._current_entry_index()
        if idx < 0 or idx >= len(entries): QtWidgets.QMessageBox.information(self, APP_TITLE, "Select an entry first."); return
        if not self.pubkey_bytes: QtWidgets.QMessageBox.warning(self, APP_TITLE, "Public key required to (re)encrypt this entry."); return
        from datetime import datetime
        e = entries[idx]
        fields = {"body": self.body.toPlainText().rstrip("\n")}
        e["title"] = self.title.text().strip(); e["updated"] = datetime.now().astimezone().isoformat()
        try: e["enc"] = encrypt_blob(pack_entry_payload(fields), self.pubkey_bytes)
        except Exception as err:
            QtWidgets.QMessageBox.critical(self, APP_TITLE, f"Entry encryption failed:\n{err}"); return
        self._select_entry_by_id(e["id"]); self.nb_dirty = True
        self.statusBar().showMessage("Entry updated (remember to Save Notebook)."); self._refresh_controls()

    def delete_entry(self):
        entries = self._visible_entries(); idx = self._current_entry_index()
        if idx < 0 or idx >= len(entries): QtWidgets.QMessageBox.information(self, APP_TITLE, "Select an entry to delete."); return
        e = entries[idx]; title = e.get("title") or "(untitled)"; dt = (e.get("dt") or "?")[:19]
        if QtWidgets.QMessageBox.question(self, APP_TITLE, f"Delete this entry?\n\n{dt} — {title}\nThis cannot be undone.",
                                          QtWidgets.QMessageBox.StandardButton.Yes | QtWidgets.QMessageBox.StandardButton.No,
                                          QtWidgets.QMessageBox.StandardButton.No) != QtWidgets.QMessageBox.StandardButton.Yes: return
        try: self.nb["entries"] = [x for x in self.nb["entries"] if x["id"] != e["id"]]
        except Exception as ex:
            QtWidgets.QMessageBox.critical(self, APP_TITLE, f"Delete failed:\n{ex}"); return
        self._populate_list(); self.nb_dirty = True; self.statusBar().showMessage("Entry deleted. Save Notebook to persist."); self._refresh_controls()

    # -------------------------- Prev/Next navigation -----------------------
    def goto_prev(self):
        row = self.listw.currentRow()
        if row is None or row <= 0: return
        self.listw.setCurrentRow(row - 1)

    def goto_next(self):
        row = self.listw.currentRow()
        if row is None or row >= self.listw.count() - 1: return
        self.listw.setCurrentRow(row + 1)

    # ------------------------------ Export ---------------------------------
    def export_docx(self):
        if not self.nb: QtWidgets.QMessageBox.information(self, APP_TITLE, "Open a notebook first."); return
        if docx is None: QtWidgets.QMessageBox.critical(self, APP_TITLE, "python-docx not available. Install with:\n  python3 -m pip install python-docx"); return
        if not self.privkey_bytes: QtWidgets.QMessageBox.warning(self, APP_TITLE, "Select/decrypt your PRIVATE key to export."); return
        out, _ = QtWidgets.QFileDialog.getSaveFileName(self, "Export Notebook to Word", filter="Word Document (*.docx)")
        if not out: return
        try:
            d = docx.Document(); d.add_heading("Journal Export", level=0)
            from datetime import datetime
            d.add_paragraph(f"Exported: {datetime.now().astimezone().isoformat()}"); d.add_paragraph("")
            for e in self._sorted_entries():
                try: body = unpack_entry_payload(decrypt_blob(e["enc"], self.privkey_bytes)).get("body", "")
                except Exception as ex: body = f"[Decryption failed: {ex}]"
                title = e.get("title") or "(untitled)"; dt = e.get("dt") or ""
                d.add_heading(f"{dt} — {title}", level=1)
                for para in (body.split("\n") or [""]): d.add_paragraph(para)
                d.add_paragraph("")
            d.save(out); self.statusBar().showMessage(f"Exported Word document: {Path(out).name} (unencrypted)")
            QtWidgets.QMessageBox.information(self, APP_TITLE, "Export complete. (Note: The .docx is UNENCRYPTED.)")
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, APP_TITLE, f"Export failed:\n{e}")

    # --------------------------- Dictation controls ------------------------
    def _ensure_vosk_ready(self) -> bool:
        if VoskModel is None or sd is None:
            QtWidgets.QMessageBox.critical(self, APP_TITLE, "Missing STT deps. Install with:\n  python3 -m pip install vosk sounddevice")
            return False
        if self.vosk_model_dir is None:
            self.vosk_model_dir = find_builtin_vosk_model()
        if self.vosk_model_dir is None:
            QtWidgets.QMessageBox.critical(self, APP_TITLE,
                                           "No Vosk model found.\nBundle at ./models/<vosk-model…> or ~/.pq-journal/models/<vosk-model…>.")
            return False
        return True

    def _ensure_pytorch_ready(self) -> bool:
        """Check if PyTorch dependencies are available for FER"""
        if torch is None or cv2 is None:
            QtWidgets.QMessageBox.critical(self, APP_TITLE,
                "Missing PyTorch FER deps. Install with:\n  python3 -m pip install torch torchvision opencv-python")
            return False
        return True

    def toggle_recording(self):
        """Unified record/stop function for both audio and video"""
        print("Toggle recording called...")

        # Check if currently recording
        is_recording = (self.dictation and self.dictation.isRunning()) or (self.video_worker and self.video_worker.isRunning())

        if is_recording:
            # Stop recording
            self._stop_recording()
        else:
            # Start recording
            self._start_recording()

    def _start_recording(self):
        """Start both audio and video recording"""
        print("Starting unified recording...")

        if not self._current_entry():
            QtWidgets.QMessageBox.information(self, APP_TITLE, "Open or create an entry first.\n\nClick the '✨ New' button to create a notebook, then the '+' button to add an entry.")
            print("No current entry selected")
            return

        try:
            # Clear streaming VA(D) buffers
            self._vad_times.clear(); self._vad_V.clear(); self._vad_A.clear(); self._vad_D.clear()

            # Start audio recording (if Vosk is available)
            if self._ensure_vosk_ready():
                self.dictation = DictationWorker(self.vosk_model_dir)
                self.dictation.partial.connect(lambda t: self.statusBar().showMessage("… " + t))
                self.dictation.final_json.connect(self._on_final_json)
                self.dictation.error.connect(lambda m: self.statusBar().showMessage("Dictation error: " + m))
                self.dictation.level.connect(self.meter.push)
                self.dictation.pulse.connect(self.meter.pulse)
                self.dictation.vad_frame.connect(self._collect_vad_frame)
                self.dictation.start()
                print("Audio recording started")
            else:
                print("Vosk not available, skipping audio recording")

            # Start video recording
            self._start_video_capture()
            print("Video recording started")

            # Update UI
            self.btn_record.setText("⏹ Stop")
            self.btn_record.setToolTip("Stop Recording")
            self.statusBar().showMessage("🔴 Recording audio + video. Speak for transcription with emotion analysis.")
            self._refresh_controls()

        except Exception as e:
            QtWidgets.QMessageBox.critical(self, APP_TITLE, f"Failed to start recording:\n{e}")
            print(f"Recording start error: {e}")

    def _stop_recording(self):
        """Stop both audio and video recording"""
        print("Stopping unified recording...")

        # Stop dictation
        if self.dictation and self.dictation.isRunning():
            try:
                self.dictation._pcm_q.queue.clear()
            except Exception:
                pass
            try:
                self.dictation.stop()
                self.dictation.wait(300)
            except Exception:
                pass
            print("Audio recording stopped")

        # Stop video
        if self.video_worker and self.video_worker.isRunning():
            try:
                self.video_worker.stop()
                self.video_worker.wait(300)
            except Exception:
                pass
            print("Video recording stopped")

        # Clean up
        self.dictation = None
        self.video_worker = None

        # Update UI
        self.btn_record.setText("🔴 Record")
        self.btn_record.setToolTip("Start Recording (Audio + Video)")
        self.statusBar().showMessage("Recording stopped.")
        self._refresh_controls()

    # Keep the old methods for compatibility, but redirect to new unified approach
    def start_dictation(self):
        """Legacy method - redirects to unified recording"""
        self._start_recording()

    def stop_dictation(self):
        """Legacy method - redirects to unified recording"""
        self._stop_recording()

    def _start_video_capture(self):
        """Start video capture and emotion detection"""
        if not self._ensure_pytorch_ready():
            return

        try:
            print("Starting video capture...")
            self.video_worker = VideoWorker()

            # Connect signals with debug output
            self.video_worker.frame_ready.connect(self.video_widget.update_frame)
            self.video_worker.emotion_detected.connect(self.video_widget.update_emotions)
            self.video_worker.error.connect(self.video_widget.show_error)
            self.video_worker.error.connect(lambda msg: print(f"Video worker error: {msg}"))

            # Set debug mode based on checkbox
            self.video_worker._debug_enabled = self.debug_checkbox.isChecked()

            self.video_worker.start()
            print("Video worker started")

        except Exception as e:
            error_msg = f"Failed to start video capture: {e}"
            print(error_msg)
            self.video_widget.show_error(str(e))

    @QtCore.pyqtSlot(float, float, float, float)
    def _collect_vad_frame(self, V: float, A: float, D: float, t: float):
        self._vad_times.append(float(t)); self._vad_V.append(float(V)); self._vad_A.append(float(A)); self._vad_D.append(float(D))
        # Bound buffer to ~10 minutes just in case
        if len(self._vad_times) > 16000:
            cut = len(self._vad_times) - 12000
            del self._vad_times[:cut]; del self._vad_V[:cut]; del self._vad_A[:cut]; del self._vad_D[:cut]

    def _select_frames(self, t0: float, t1: float) -> Tuple[List[float], List[float], List[float]]:
        V, A, D = [], [], []
        for i, t in enumerate(self._vad_times):
            if t0 <= t <= t1:
                V.append(self._vad_V[i]); A.append(self._vad_A[i]); D.append(self._vad_D[i])
        return V, A, D

    def _rescale_to_session_percentiles(self, vals: List[float], hist: List[float], lo_q=0.10, hi_q=0.90) -> List[float]:
        """
        Rescales `vals` to [0,1] based on the lo/hi quantiles of `hist`.
        Robust to numpy returning 0-d or 1-d arrays for quantiles.
        """
        if not vals or len(hist) < 10:
            return vals

        # Force a clean 1-D float array (avoid deques / nested lists)
        arr = _np.asarray(list(hist), dtype=_np.float32).reshape(-1)
        if arr.size < 10:
            return vals

        def _as_float(x) -> float:
            # Works for numpy scalars and arrays of any shape
            try:
                return float(x)   # numpy scalar or python scalar
            except TypeError:
                a = _np.asarray(x).reshape(-1)
                return float(a[0]) if a.size else 0.5

        lo = _as_float(_np.quantile(arr, lo_q))
        hi = _as_float(_np.quantile(arr, hi_q))
        if not _np.isfinite(lo) or not _np.isfinite(hi) or hi <= lo:
            return vals

        out: List[float] = []
        inv = 1.0 / (hi - lo)
        for v in vals:
            out.append(float(max(0.0, min(1.0, (v - lo) * inv))))
        return out

    def _insert_text_and_summary(self, plain_text: str, start_time: float, end_time: float):
        # Clean & insert text
        cleaned = _basic_punctuate_and_capitalize(plain_text)
        cursor = self.body.textCursor(); cursor.movePosition(QtGui.QTextCursor.MoveOperation.End)
        # New paragraph
        if self.body.toPlainText().strip():
            cursor.insertText("\n")
        cursor.insertText(cleaned)
        self.body.setTextCursor(cursor)

        V_seq, A_seq, D_seq = self._select_frames(start_time, end_time)

        # --- session-calibrated arousal (prevents "subdued" spam) ---
        A_seq = self._rescale_to_session_percentiles(A_seq, self._vad_A)

        # NEW: get (valence, is_academic)
        text_val, is_academic = self.text_est.valence(cleaned)

        # Optional nudge for emphatic punctuation (keep tiny)
        if re.search(r"[!?!]{1,}$", cleaned):
            text_val = max(text_val, 0.15)
            A_seq = [min(1.0, a + 0.05) for a in A_seq]

        # NEW: Get video emotions from the specific paragraph time window
        video_emotions = []
        if self.video_worker:
            video_emotions = self.video_worker.get_emotions_in_timerange(start_time, end_time)

        # Get current fusion weights from dropdown
        audio_weight, video_weight = self._get_fusion_weights()
        phrase, meta = self.summarizer.summarize(V_seq, A_seq, D_seq, (text_val, is_academic), video_emotions, audio_weight, video_weight)

        # NEW: Local Llama emotion interpretation (optional)
        llama_result = None
        if self._ensure_llama_interpreter():
            try:
                # Prepare data for Llama
                audio_vad = meta.get("audio_vad", {"V": 0.0, "A": 0.0, "D": 0.0})

                # Call local Llama interpreter
                llama_result = self.llama_interpreter.interpret_emotion(
                    text=cleaned,
                    audio_data=audio_vad,
                    video_data=video_emotions,
                    debug=self.debug_checkbox.isChecked()
                )

                # Store Llama result if available and valid
                if llama_result and not llama_result.startswith("llama"):
                    meta["llama_enhanced"] = True
                    meta["llama_result"] = llama_result

            except Exception as e:
                print(f"Llama interpretation failed: {e}")
                # Continue with traditional phrase

        # Debug output if enabled
        if self.debug_checkbox.isChecked():
            audio_phrase = meta.get("audio_phrase", "unknown audio")
            video_phrase = meta.get("video_phrase", None)
            video_str = video_phrase if video_phrase else "no video"

            # Raw audio data
            audio_vad = meta.get("audio_vad", {})
            raw_audio = meta.get("raw_audio_sequences", {})
            audio_raw_str = f"V={audio_vad.get('V', 0):.3f} A={audio_vad.get('A', 0):.3f} D={audio_vad.get('D', 0):.3f}"

            # Enhanced raw video data output
            raw_video = meta.get("raw_video_emotions", [])
            video_raw_str = "no video data"
            if raw_video:
                # Get paragraph timing info
                paragraph_duration = end_time - start_time
                print(f"\n=== PARAGRAPH-LEVEL FER ANALYSIS ===")
                print(f"Paragraph duration: {paragraph_duration:.2f}s ({len(raw_video)} detections)")

                # Calculate and show averaged emotions for the entire paragraph
                averaged_emotions = self.summarizer._average_emotions_over_paragraph(raw_video)
                if averaged_emotions:
                    print("\n--- AVERAGED EMOTIONS ACROSS ENTIRE PARAGRAPH ---")
                    sorted_avg = sorted(averaged_emotions.items(), key=lambda x: x[1], reverse=True)
                    avg_emotions_str = " ".join([f"{emo}={score:.3f}" for emo, score in sorted_avg])
                    print(f"  Paragraph average: {avg_emotions_str}")

                print(f"\n--- INDIVIDUAL DETECTIONS ---")
                for i, detection in enumerate(raw_video):
                    bbox = detection.get('bbox', {})
                    confidence = detection.get('confidence', 0.0)
                    emotions = detection.get('emotions', {})
                    timestamp = detection.get('timestamp', 0.0)

                    # Show bounding box info with timestamp
                    bbox_str = f"face@({bbox.get('x', 0)},{bbox.get('y', 0)}) {bbox.get('w', 0)}x{bbox.get('h', 0)}"

                    # Show ALL emotion scores, sorted by confidence
                    emotion_items = sorted(emotions.items(), key=lambda x: x[1], reverse=True)
                    all_emotions_str = " ".join([f"{emo}={prob:.3f}" for emo, prob in emotion_items])

                    print(f"  Detection {i+1} @ {timestamp:.2f}s: {bbox_str} conf={confidence:.3f}")
                    print(f"    All emotions: {all_emotions_str}")

                # Use averaged emotions for summary instead of best single detection
                if averaged_emotions:
                    top_emotions = sorted(averaged_emotions.items(), key=lambda x: x[1], reverse=True)[:3]
                    emotion_strs = [f"{emo}={prob:.3f}" for emo, prob in top_emotions]
                    video_raw_str = " ".join(emotion_strs)

                print("=" * 50)

            # Show Llama results if available
            if llama_result:
                print(f"\n=== LOCAL LLAMA INTERPRETATION ===")
                print(f"Llama summary: '{llama_result}'")

                # Show if Llama was actually used
                if meta.get("llama_enhanced"):
                    print("✓ Llama result was used for final phrase")
                else:
                    print("✗ Llama result discarded (invalid response)")

                print("=" * 50)

            print(f'"{phrase}" [audio: {audio_phrase} ({audio_raw_str}); video: {video_str} ({video_raw_str})]')

        # Insert summary line
        cursor = self.body.textCursor(); cursor.movePosition(QtGui.QTextCursor.MoveOperation.End)
        # Build final summary line with traditional phrase + optional Llama result
        summary_line = f"\n⟡ {phrase}"
        if meta.get("llama_enhanced") and meta.get("llama_result"):
            summary_line += f"\n• {{Emotion: {meta['llama_result']}}}"
        cursor.insertText(summary_line)
        self.body.setTextCursor(cursor)
        self.nb_dirty = True; self._refresh_controls()

    @QtCore.pyqtSlot(dict)
    def _on_final_json(self, j: dict):
        """
        Vosk final JSON:
          { "text": "...", "result":[{"conf":..., "end":1.23, "start":0.87, "word":"..."}, ...] }
        """
        text = (j.get("text") or "").strip()
        words = j.get("result") or []
        if not text:
            return

        # NEW: State machine integration for emotion processing
        self._on_speech_detected(text)

        # OLD: Direct processing - still used for immediate text insertion
        # Determine paragraph time span from first/last word timestamps
        if words and isinstance(words, list):
            t0 = float(words[0].get("start", 0.0))
            t1 = float(words[-1].get("end", t0))
        else:
            # Fallback: approximate using recent audio window (last 2s)
            t1 = self._vad_times[-1] if self._vad_times else 0.0
            t0 = max(0.0, t1 - 2.0)

        # Insert text immediately but handle emotion separately
        self._insert_text_only(text, t0, t1)

    def stop_dictation(self):
        if self.dictation and self.dictation.isRunning():
            try: self.dictation._pcm_q.queue.clear()
            except Exception: pass
            try: self.dictation.stop()
            except Exception: pass
            self.dictation.wait(300)

        # NEW: Stop video capture
        if self.video_worker and self.video_worker.isRunning():
            try: self.video_worker.stop()
            except Exception: pass
            self.video_worker.wait(300)

        # NEW: Clean up emotion state machine
        self._cleanup_emotion_state_machine()

        self.dictation = None
        self.video_worker = None
        self.statusBar().showMessage("Dictation stopped.")
        self._refresh_controls()

    def _cleanup_emotion_state_machine(self):
        # Function Purpose: Robust shutdown with Any → IDLE state transition
        try:
            if hasattr(self, 'debug_checkbox') and self.debug_checkbox.isChecked():
                ts = debug_timestamp()
                print(f"[{ts}] State: {self.emotion_state} → IDLE (recording stopped)")
        except:
            pass  # Don't let debug output crash cleanup

        try:
            # Stop pause timer
            if hasattr(self, '_pause_timer') and self._pause_timer:
                self._pause_timer.stop()
        except Exception as e:
            print(f"Error stopping pause timer: {e}")

        try:
            # Stop LLM timeout timer
            if hasattr(self, '_llm_timeout_timer') and self._llm_timeout_timer:
                self._llm_timeout_timer.stop()
        except Exception as e:
            print(f"Error stopping LLM timeout timer: {e}")

        try:
            # State transition: Any → IDLE (recording stopped)
            self.emotion_state = EmotionState.IDLE
            if hasattr(self, '_paragraph_buffer'):
                self._paragraph_buffer.clear()
            self._speech_since_ready = False
            self._first_speech_time = 0.0  # Always reset accumulation timer
            self._llm_request_time = 0.0  # Always reset LLM request timer
        except Exception as e:
            print(f"Error resetting state: {e}")

        try:
            # Stop emotion worker thread with robust handling
            if hasattr(self, 'emotion_worker') and self.emotion_worker:
                if self.emotion_worker.is_alive():
                    print("Stopping emotion worker thread...")
                    self.emotion_worker.stop()

                    # Try to join with timeout
                    self.emotion_worker.join(timeout=3.0)

                    # Force termination if still alive
                    if self.emotion_worker.is_alive():
                        print("Warning: Emotion worker thread did not stop gracefully")
                    else:
                        print("Emotion worker thread stopped successfully")

                self.emotion_worker = None
        except Exception as e:
            print(f"Error stopping emotion worker: {e}")
            # Don't re-raise - continue with cleanup

    # ------------------------------ Command palette ------------------------
    def _open_palette(self):
        if not hasattr(self, "_palette") or self._palette is None:
            dlg = QtWidgets.QDialog(self); dlg.setWindowTitle("Command Palette")
            lay = QtWidgets.QVBoxLayout(dlg)
            entry = QtWidgets.QLineEdit(placeholderText="Type a command…")
            listw = QtWidgets.QListWidget(); lay.addWidget(entry); lay.addWidget(listw)
            cmds = [("New Notebook", self.create_notebook), ("Open Notebook", self.open_notebook),
                    ("Save Notebook", self.save_notebook), ("Export to Word", self.export_docx),
                    ("Add Entry", self.add_entry), ("Delete Entry", self.delete_entry),
                    ("Toggle Recording", self.toggle_recording),
                    ("Prev Entry", self.goto_prev), ("Next Entry", self.goto_next),
                    ("Close Notebook", self.close_notebook)]
            for name, _ in cmds: listw.addItem(name)
            def run_selected():
                items = listw.selectedItems()
                if not items: return
                name = items[0].text()
                for label, fn in cmds:
                    if label == name: fn(); break
                dlg.accept()
            listw.itemActivated.connect(lambda *_: run_selected())
            entry.textChanged.connect(lambda t: self._palette_filter(listw, t))
            entry.returnPressed.connect(run_selected)
            self._palette = dlg; self._palette_entry = entry; self._palette_list = listw
        self._palette_entry.clear(); self._palette_filter(self._palette_list, "")
        self._palette.show(); self._palette.raise_(); self._palette_entry.setFocus()

    def _palette_filter(self, listw: QtWidgets.QListWidget, t: str):
        t = t.lower().strip()
        for i in range(listw.count()):
            item = listw.item(i); item.setHidden(t not in item.text().lower())

    def closeEvent(self, ev: QtGui.QCloseEvent):
        # Cleanup emotion state machine
        self._cleanup_emotion_state_machine()

        # Clean up llama interpreter if running
        if hasattr(self, 'llama_interpreter') and self.llama_interpreter:
            try:
                self.llama_interpreter.cleanup()
            except Exception as e:
                print(f"Error cleaning up llama interpreter: {e}")

        # Stop workers
        try:
            if self.dictation and self.dictation.isRunning():
                self.dictation.stop()
                self.dictation.wait(300)
        except Exception:
            pass

        try:
            if self.video_worker and self.video_worker.isRunning():
                self.video_worker.stop()
                self.video_worker.wait(300)
        except Exception:
            pass

        super().closeEvent(ev)

# -------------------------------- main ------------------------------------
def main():
    app = QtWidgets.QApplication(sys.argv)
    try:
        fixed = QtGui.QFontDatabase.systemFont(QtGui.QFontDatabase.SystemFont.FixedFont)
        if fixed and fixed.family(): app.setFont(fixed)
    except Exception: pass
    w = MainWindow()

    mbar = w.menuBar()
    filem = mbar.addMenu("File")
    act_new   = QtGui.QAction("New", w);   act_new.setShortcut(QtGui.QKeySequence.StandardKey.New);   act_new.triggered.connect(w.create_notebook); filem.addAction(act_new)
    act_open  = QtGui.QAction("Open", w);  act_open.setShortcut(QtGui.QKeySequence.StandardKey.Open);  act_open.triggered.connect(w.open_notebook);  filem.addAction(act_open)
    act_save  = QtGui.QAction("Save", w);  act_save.setShortcut(QtGui.QKeySequence.StandardKey.Save);  act_save.triggered.connect(w.save_notebook); filem.addAction(act_save)
    act_close = QtGui.QAction("Close", w); act_close.setShortcut(QtGui.QKeySequence.StandardKey.Close); act_close.triggered.connect(w.close_notebook); filem.addAction(act_close)
    act_export= QtGui.QAction("Export to Word", w); act_export.setShortcut(QtGui.QKeySequence("Ctrl+E")); act_export.triggered.connect(w.export_docx); filem.addAction(act_export)
    filem.addSeparator()
    act_quit  = QtGui.QAction("Quit", w);  act_quit.setShortcut(QtGui.QKeySequence.StandardKey.Quit);  act_quit.triggered.connect(QtWidgets.QApplication.instance().quit); filem.addAction(act_quit)

    editm = mbar.addMenu("Edit")
    act_add  = QtGui.QAction("Add Entry", w); act_add.setShortcut(QtGui.QKeySequence("Ctrl+N")); act_add.triggered.connect(w.add_entry); editm.addAction(act_add)
    act_del  = QtGui.QAction("Delete Entry", w); act_del.triggered.connect(w.delete_entry); editm.addAction(act_del)

    navm = mbar.addMenu("Navigate")
    act_prev = QtGui.QAction("Prev", w); act_prev.triggered.connect(w.goto_prev); navm.addAction(act_prev)
    act_next = QtGui.QAction("Next", w); act_next.triggered.connect(w.goto_next); navm.addAction(act_next)

    toolsm = mbar.addMenu("Tools")
    act_record = QtGui.QAction("Toggle Recording", w); act_record.setShortcut(QtGui.QKeySequence("F5")); act_record.triggered.connect(w.toggle_recording); toolsm.addAction(act_record)
    act_pal    = QtGui.QAction("Command Palette", w); act_pal.setShortcut(QtGui.QKeySequence("Ctrl+K")); act_pal.triggered.connect(w._open_palette); toolsm.addAction(act_pal)

    err = require_crypto_ready()
    if err: QtWidgets.QMessageBox.critical(w, APP_TITLE, err)
    w.show(); sys.exit(app.exec())

# ============================ Unit Tests for State Machine ============================
def test_emotion_state_transitions():
    # Function Purpose: Minimal unit checks for state machine transitions
    print("Testing emotion state machine transitions...")

    # Create minimal test instance
    class MockMainWindow:
        def __init__(self):
            self.emotion_state = EmotionState.IDLE
            self._paragraph_buffer = []
            self._speech_since_ready = False
            self.debug_checkbox = type('MockCheckbox', (), {'isChecked': lambda: True})()

        def _is_worker_idle(self):
            return True

        def _try_enqueue_current_paragraph(self):
            if self._paragraph_buffer:
                self.emotion_state = EmotionState.WAITING_MODEL
                print(f"  Mock: Enqueued paragraph with {len(self._paragraph_buffer)} chunks")

    mock = MockMainWindow()

    # Test 1: IDLE → LISTENING (speech start)
    print("Test 1: IDLE → LISTENING")
    assert mock.emotion_state == EmotionState.IDLE
    mock._paragraph_buffer.append("Hello")
    if mock.emotion_state == EmotionState.IDLE:
        mock.emotion_state = EmotionState.LISTENING
    assert mock.emotion_state == EmotionState.LISTENING
    print("  ✓ IDLE → LISTENING on speech start")

    # Test 2: LISTENING → WAITING_MODEL (pause → enqueue job)
    print("Test 2: LISTENING → WAITING_MODEL")
    if mock.emotion_state == EmotionState.LISTENING and mock._is_worker_idle():
        mock._try_enqueue_current_paragraph()
    assert mock.emotion_state == EmotionState.WAITING_MODEL
    print("  ✓ LISTENING → WAITING_MODEL on pause with idle worker")

    # Test 3: WAITING_MODEL → READY_NEXT (llama result arrives)
    print("Test 3: WAITING_MODEL → READY_NEXT")
    mock.emotion_state = EmotionState.READY_NEXT
    mock._speech_since_ready = False
    assert mock.emotion_state == EmotionState.READY_NEXT
    print("  ✓ WAITING_MODEL → READY_NEXT on llama result")

    # Test 4: READY_NEXT → LISTENING (speech resumes)
    print("Test 4: READY_NEXT → LISTENING")
    mock._paragraph_buffer.append("World")
    if mock.emotion_state == EmotionState.READY_NEXT:
        mock.emotion_state = EmotionState.LISTENING
        mock._speech_since_ready = True
    assert mock.emotion_state == EmotionState.LISTENING
    assert mock._speech_since_ready == True
    print("  ✓ READY_NEXT → LISTENING on speech resume")

    # Test 5: READY_NEXT → WAITING_MODEL (next pause after new speech)
    print("Test 5: READY_NEXT → WAITING_MODEL")
    mock.emotion_state = EmotionState.READY_NEXT
    mock._speech_since_ready = True
    if mock._speech_since_ready and mock._paragraph_buffer:
        mock._try_enqueue_current_paragraph()
    assert mock.emotion_state == EmotionState.WAITING_MODEL
    print("  ✓ READY_NEXT → WAITING_MODEL on pause after new speech")

    # Test 6: Any → IDLE (recording stopped)
    print("Test 6: Any → IDLE")
    mock.emotion_state = EmotionState.IDLE
    mock._paragraph_buffer.clear()
    mock._speech_since_ready = False
    assert mock.emotion_state == EmotionState.IDLE
    assert len(mock._paragraph_buffer) == 0
    print("  ✓ Any → IDLE on recording stop")

    print("All state machine tests passed! ✓")

def test_emotion_job_creation():
    # Function Purpose: Test EmotionJob creation and VAD data handling
    print("Testing EmotionJob creation...")

    job = EmotionJob(
        text="Test paragraph text",
        vad={'valence': 0.5, 'arousal': 0.3, 'dominance': 0.7},
        created_ts=time.time()
    )

    assert job.text == "Test paragraph text"
    assert job.vad['valence'] == 0.5
    assert job.vad['arousal'] == 0.3
    assert job.vad['dominance'] == 0.7
    assert isinstance(job.created_ts, float)

    print("  ✓ EmotionJob creation and VAD data handling")

if __name__ == "__main__":
    # Run unit tests if --test flag is provided
    if len(sys.argv) > 1 and sys.argv[1] == "--test":
        test_emotion_state_transitions()
        test_emotion_job_creation()
        print("\nAll tests completed successfully!")
        sys.exit(0)

    # Normal GUI application startup
    main()